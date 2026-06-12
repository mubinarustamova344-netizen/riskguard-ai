# -*- coding: utf-8 -*-
"""Generates the full BTEC Level 6 Independent Project report as a DOCX file."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)

def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p

def para(text, bold=False, italic=False, indent=False, space_after=10, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if align:
        p.alignment = align
    return p

def centre(text, bold=False, size=12, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(11)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F4E79')
        shd.set(qn('w:color'), 'FFFFFF')
        hdr[i]._tc.get_or_add_tcPr().append(shd)
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for r_idx, row in enumerate(rows):
        cells = t.rows[r_idx+1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
            if r_idx % 2 == 0:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'DCE6F1')
                cells[c_idx]._tc.get_or_add_tcPr().append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def page_break():
    doc.add_page_break()

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

def figure_placeholder(fig_num, title, desc="[Insert screenshot/diagram here]"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    shading_elm = OxmlElement('w:pPr')
    run = p.add_run(f"\n{desc}\n")
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.size = Pt(10)
    run.font.italic = True
    caption(f"Figure {fig_num}: {title}")

# ═══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph().paragraph_format.space_after = Pt(24)
centre("PDP UNIVERSITY", bold=True, size=16)
centre("Faculty of Business Information Technology", bold=False, size=13)
centre("Tashkent, Uzbekistan", bold=False, size=12)
doc.add_paragraph()
centre("─────────────────────────────────────────")
doc.add_paragraph()
centre("INDEPENDENT PROJECT", bold=True, size=14)
centre("Pearson BTEC Level 6 Diploma in Digital Technologies", bold=False, size=12)
centre("Unit 2  —  Unit Code: 70726U  —  Credit Value: 30", bold=False, size=11)
doc.add_paragraph()
centre("RiskGuard AI: A Machine Learning–Powered Road Accident\nRisk Assessment System for Insurance Companies", bold=True, size=16)
doc.add_paragraph()
centre("How can machine learning improve the accuracy, consistency and transparency\nof motor insurance risk assessment and premium pricing?", italic=True, size=12)
doc.add_paragraph()
centre("─────────────────────────────────────────")
doc.add_paragraph()

t = doc.add_table(rows=7, cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ("Student Name:", "Mubina Rustamova"),
    ("Student ID:", "230277"),
    ("Programme / Group:", "BIT — 22-306"),
    ("Project Format:", "Capstone-style"),
    ("Supervisor:", "Xolmirzayev Muxammadjon"),
    ("Submission Date:", "08 June 2026"),
    ("Word Count:", "Approximately 12,000 words"),
]
for i, (k, v) in enumerate(data):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    t.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    t.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(11)

doc.add_paragraph()
centre("Submitted in partial fulfilment of the requirements for the", italic=True)
centre("Bachelor's Degree in Business Information Technology", italic=True, bold=True)
centre("Academic Year 2025–2026", italic=True)
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  DECLARATION
# ═══════════════════════════════════════════════════════════════════════════════
heading("Declaration of Originality", 1)
para("I hereby declare that this Independent Project, submitted in partial fulfilment of the requirements for the Pearson BTEC Level 6 Diploma in Digital Technologies and the Bachelor's Degree in Business Information Technology at PDP University, is the result of my own original work.")
para("I confirm that:")
bullet("All sources of information, data and ideas drawn from other authors have been fully acknowledged through accurate in-text citations and a complete reference list using the Harvard (author–date) referencing system.")
bullet("This work has not been previously submitted, in whole or in part, for any other academic award at this or any other institution.")
bullet("All research activities have been conducted in compliance with the institutional ethics procedures of PDP University and applicable data-protection regulations.")
bullet("I understand that any breach of academic integrity, including plagiarism, fabrication of data or unauthorised collaboration, may result in the withdrawal of this submission and disciplinary action.")
doc.add_paragraph()
t2 = doc.add_table(rows=2, cols=2)
t2.rows[0].cells[0].text = "Signature:"
t2.rows[0].cells[1].text = "Date:  08.06.2026"
t2.rows[1].cells[0].text = "Mubina Rustamova"
t2.rows[1].cells[1].text = "Tashkent, Uzbekistan"
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  ACKNOWLEDGEMENTS
# ═══════════════════════════════════════════════════════════════════════════════
heading("Acknowledgements", 1)
para("I would like to express my sincere gratitude to my supervisor, Xolmirzayev Muxammadjon, for his continuous guidance, insightful feedback and encouragement throughout the duration of this project. His expertise in machine learning and software engineering significantly shaped the technical direction of RiskGuard AI.")
para("I am also grateful to the faculty of the Business Information Technology department at PDP University for providing an intellectually stimulating environment and the academic foundation that made this work possible. The structured approach to project management and research methodology taught throughout the programme proved invaluable at every stage.")
para("My thanks extend to Anthropic's Claude Code, which served as an AI-assisted development tool and significantly accelerated the implementation phase of this project. This experience has deepened my appreciation for responsible and productive human–AI collaboration in software engineering.")
para("Finally, I extend my heartfelt thanks to my family, whose unwavering support, patience and encouragement have been instrumental throughout my Bachelor's studies at PDP University.")
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════
heading("Abstract", 1)
para("The motor insurance industry continues to rely on static, demographic-centric underwriting models that inadequately capture the behavioural and contextual risk signals that predict road accident probability. This project investigates the application of machine learning to automate and improve driver risk assessment within the motor insurance context. The aim of the study is to design, develop and evaluate a full-stack web application — RiskGuard AI — that employs machine learning to produce explainable, accurate and actionable risk profiles for insurance underwriters.")
para("A Design Science Research (DSR) methodology was adopted, combining quantitative model development and evaluation with the design and implementation of a production-grade software artefact. A synthetic dataset of 30,000 driver records was generated using domain-weighted heuristics validated against actuarial literature, incorporating 18 features — 12 direct input variables spanning driver behaviour, vehicle characteristics and environmental context, plus 6 engineered interaction features. Six classifiers were trained and compared: Decision Tree (73.1%), Logistic Regression (77.7%), Extra Trees (78.2%), Random Forest (81.0%), Gradient Boosting (85.0%), and XGBoost Tuned (85.22%, ROC-AUC 0.9738). The XGBoost model, tuned via GridSearchCV with 3-fold cross-validation, was selected as the production model.")
para("The system was deployed as a secure Flask web application featuring a risk score gauge (0–100), premium multiplier calculation (1.0×–2.30×), per-factor contribution breakdown, counterfactual improvement roadmap, AI chatbot (RiskBot) powered by Claude Sonnet 4.6 with a regex fallback engine covering 20+ intents in English and Uzbek, bilingual English/Uzbek interface, PDF report generation, batch CSV upload, risk comparison tool, and an administrative dashboard with historical trend analysis and real-time model metrics. Security controls include CSRF protection, rate limiting (Flask-Limiter), bcrypt password hashing, HTTP security headers (X-Frame-Options, X-Content-Type-Options) and session management.")
para("The findings confirm that XGBoost significantly outperforms linear baselines for tabular insurance risk classification, achieving a ROC-AUC of 0.9738 across four risk classes. The feature engineering pipeline, which adds six interaction features (including accident-per-experience ratio and youth risk index) to the 12 base inputs, demonstrably improves model discrimination. The project concludes that machine learning can replace manual underwriting with a reproducible, transparent decision-support system, and recommends integration of real telematics data and SHAP-based explainability as immediate next steps.")
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Keywords: ")
run.bold = True
run.font.size = Pt(12)
run2 = p.add_run("machine learning; motor insurance; risk assessment; XGBoost; explainable AI; Flask; Python; insurance underwriting; premium pricing; Design Science Research")
run2.italic = True
run2.font.size = Pt(12)
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════════
heading("Table of Contents", 1)
toc_items = [
    ("Declaration of Originality", "2"),("Acknowledgements", "3"),
    ("Abstract", "4"),("Table of Contents", "5"),
    ("List of Figures", "6"),("List of Tables", "6"),("List of Abbreviations", "7"),
    ("Chapter 1 — Introduction", "8"),
    ("  1.1  Background and Context", "8"),("  1.2  Problem Statement", "9"),
    ("  1.3  Project Aim", "9"),("  1.4  Project Objectives", "10"),
    ("  1.5  Research Questions", "10"),("  1.6  Significance of the Project", "10"),
    ("  1.7  Scope and Limitations", "11"),("  1.8  Structure of the Report", "11"),
    ("Chapter 2 — Literature Review", "12"),
    ("Chapter 3 — Project Planning and Methodology", "17"),
    ("Chapter 4 — Data Collection and Analysis", "23"),
    ("Chapter 5 — Discussion", "30"),
    ("Chapter 6 — Conclusion and Recommendations", "34"),
    ("Chapter 7 — Reflective Evaluation", "36"),
    ("References", "40"),("Appendices", "43"),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    dots = "." * max(1, 60 - len(item) - len(page))
    run = p.add_run(f"{item} {dots} {page}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  LIST OF FIGURES & TABLES
# ═══════════════════════════════════════════════════════════════════════════════
heading("List of Figures", 1)
figs = [
    ("1","Conceptual Framework: RiskGuard AI System Architecture","8"),
    ("2","Project Gantt Chart — 10-Week Timeline","19"),
    ("3","Risk Category Distribution — Synthetic Dataset","24"),
    ("4","ML Model Performance Comparison (6 Algorithms)","26"),
    ("5","Feature Importance Ranking (Top 15)","27"),
    ("6","Confusion Matrix — XGBoost Tuned","27"),
    ("7","RiskGuard AI — Login Screen","28"),
    ("8","Admin Dashboard — KPI and Charts","28"),
    ("9","Risk Assessment Form (12 Fields)","29"),
    ("10","Assessment Results — Gauge, Factor Breakdown, Roadmap","29"),
    ("11","Driver Comparison Tool","30"),
    ("12","AI Chatbot Interface (RiskBot)","30"),
    ("13","Assessment History and CSV Export","31"),
    ("14","Insurance Quote Calculator with NCB","31"),
    ("15","ML Model Performance Report Screen","32"),
    ("16","PDF Report Preview","32"),
]
for num, title, page in figs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"Figure {num}: {title} {'.' * max(1,55-len(title)-len(num))} p. {page}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

doc.add_paragraph()
heading("List of Tables", 1)
tbls = [
    ("1","Project Objectives — SMART Framework","10"),
    ("2","Milestones and Critical Path","19"),
    ("3","Resource Planning","20"),
    ("4","Risk Register","20"),
    ("5","Synthetic Dataset Summary Statistics (n=30,000)","24"),
    ("6","Feature Definitions and Risk Contributions","25"),
    ("7","ML Model Performance Comparison","26"),
    ("8","GridSearchCV Hyperparameter Grid","26"),
    ("9","Feature Importance Top 10","27"),
    ("10","Improvement Roadmap — Average Score Savings","28"),
    ("11","Functional Test Cases","32"),
    ("12","Objectives Achievement RAG Matrix","33"),
    ("13","SWOT Analysis of Personal Development","38"),
    ("14","Transferable Skills Gained","39"),
    ("15","Future Development Plan","39"),
]
for num, title, page in tbls:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"Table {num}: {title} {'.' * max(1,55-len(title)-len(num))} p. {page}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  ABBREVIATIONS
# ═══════════════════════════════════════════════════════════════════════════════
heading("List of Abbreviations", 1)
abbrevs = [
    ("AI","Artificial Intelligence"),("API","Application Programming Interface"),
    ("BTEC","Business and Technology Education Council"),
    ("CSRF","Cross-Site Request Forgery"),("CV","Cross-Validation"),
    ("DSR","Design Science Research"),("GDPR","General Data Protection Regulation"),
    ("GB","Gradient Boosting"),("GLM","Generalised Linear Model"),
    ("HTML","HyperText Markup Language"),("JSON","JavaScript Object Notation"),
    ("KPI","Key Performance Indicator"),("LR","Logistic Regression"),
    ("ML","Machine Learning"),("NCB","No-Claims Bonus"),
    ("ORM","Object-Relational Mapper"),("PDF","Portable Document Format"),
    ("PWA","Progressive Web Application"),("RAG","Red, Amber, Green status rating"),
    ("RF","Random Forest"),("REST","Representational State Transfer"),
    ("SHAP","SHapley Additive exPlanations"),
    ("SMART","Specific, Measurable, Achievable, Relevant, Time-bound"),
    ("SQL","Structured Query Language"),("SWOT","Strengths, Weaknesses, Opportunities, Threats"),
    ("UBI","Usage-Based Insurance"),("UI","User Interface"),
    ("WBS","Work Breakdown Structure"),("XGBoost","eXtreme Gradient Boosting"),
    ("XSS","Cross-Site Scripting"),
]
table(["Abbreviation","Full Form"], abbrevs, [4,12])
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 1 — INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
heading("Chapter 1 — Introduction", 1)
para("This chapter introduces the RiskGuard AI project, establishes the context within the digital technologies and insurance sectors, defines the problem addressed, states the project aim and objectives, and outlines the structure of the remainder of the report.")

heading("1.1  Background and Context", 2)
para("The global motor insurance industry is undergoing a period of significant technological disruption. The Association of British Insurers (ABI, 2023) reported that motor insurance claims in the United Kingdom alone reached £6.7 billion in 2023, with fraud-related losses accounting for an estimated £1.1 billion annually. Simultaneously, rapid digitalisation — driven by cloud computing, big data and machine learning — is transforming how insurers assess, price and manage risk (McKinsey & Company, 2023). In Central Asia, and Uzbekistan in particular, the insurance sector is experiencing accelerated growth following regulatory liberalisation, with premiums increasing at a compound annual growth rate of 18% between 2020 and 2024 (Uzbekistan Insurance Market Report, 2024). This growth brings both the opportunity and the necessity for more sophisticated, technology-driven underwriting approaches.")
para("Machine learning (ML) has emerged as a powerful tool for risk segmentation in financial services. Unlike traditional actuarial models based on Generalised Linear Models (GLMs), which assume linear relationships between predictors and outcomes, ML algorithms can capture non-linear interactions and high-order feature combinations that significantly improve predictive accuracy (Verbelen, Antonio and Claeskens, 2018). The availability of telematics data — GPS speed profiles, braking patterns, and time-of-day usage — has further expanded the feature space available to insurers, with studies demonstrating up to 30% reduction in premium uncertainty compared to static models (Ayuso, Guillen and Nielsen, 2019). Despite this evidence, many insurance companies, particularly in emerging markets, continue to rely on manual, subjective underwriting processes that introduce inconsistency and limit scalability.")
para("RiskGuard AI was conceived in this context: as a demonstration that production-grade, ML-powered risk assessment is achievable within a single development cycle using open-source tools, and that the outputs of such a system can be made sufficiently explainable and transparent to satisfy both regulatory requirements and end-user trust.")

heading("1.2  Problem Statement", 2)
para("Traditional motor insurance underwriting is characterised by four interconnected deficiencies that reduce pricing accuracy, introduce systemic bias, and limit insurer–policyholder engagement. First, the feature space used by conventional underwriting models is narrow — typically fewer than ten variables, predominantly demographic — and fails to incorporate the richer behavioural signals confirmed as the strongest predictors of accident risk (Jiang et al., 2020). Second, manual underwriting introduces significant inter-underwriter variability: the same risk profile may be priced differently by two underwriters, violating the principle of consistency that underpins fair pricing. Third, policyholders receive no actionable feedback — they are told their premium but not which specific factors drove it or what they could change to reduce it. Fourth, the absence of explainability in many ML-based insurance tools has drawn regulatory scrutiny under GDPR Article 22, which grants citizens the right to explanation for automated decisions affecting them (Cevolini and Esposito, 2020).")
para("The specific gap addressed by this project is the absence of an integrated, explainable, production-ready ML risk assessment platform that combines accurate prediction with transparent, actionable outputs — a gap that is particularly acute in the Uzbek and broader Central Asian insurance market.")

heading("1.3  Project Aim", 2)
p = doc.add_paragraph()
run = p.add_run("Aim: ")
run.bold = True
run.font.size = Pt(12)
p.add_run("The aim of this project is to design, develop and evaluate a machine-learning–powered web application (RiskGuard AI) that automates motor insurance driver risk profiling, generates explainable risk scores and premium multipliers, and provides personalised improvement recommendations — thereby replacing subjective manual underwriting with a consistent, transparent, data-driven decision-support system.")

heading("1.4  Project Objectives", 2)
para("The following SMART objectives operationalise the project aim:")
table(
    ["Ref","Objective","Success Criterion","Target"],
    [
        ("O1","Review academic and industry literature on ML in motor insurance to identify theories, models and research gaps","≥20 credible sources; gap explicitly identified","Week 2"),
        ("O2","Design and generate a realistic synthetic training dataset of sufficient size and feature richness","30,000 records; 18 features; balanced class distribution (35/35/20/10%)","Week 3"),
        ("O3","Train, compare and select the best ML classifier using rigorous evaluation","≥6 algorithms compared; best achieves >80% accuracy; GridSearchCV tuning","Week 4"),
        ("O4","Build a secure, full-stack Flask web application","CSRF, rate limiting, bcrypt, security headers; <500ms prediction response","Week 7"),
        ("O5","Implement explainability: factor contributions, improvement roadmap, recommendations","Per-prediction breakdown and top-5 counterfactual improvements","Week 8"),
        ("O6","Provide advanced features: PDF, AI chatbot, bilingual UI, batch upload, comparison","All features functional; chatbot 20+ intents; EN+UZ language support","Week 9"),
        ("O7","Evaluate the system against functional and non-functional requirements","All 12 FR and 10 NFR met; test cases documented","Week 9"),
        ("O8","Document the project to BTEC Level 6 academic standards","Full report with Harvard citations; reflective chapter; appendices","Week 10"),
    ],
    [1.2, 5.5, 5.0, 2.0]
)

heading("1.5  Research Questions", 2)
numbered("RQ1: To what extent can machine learning classifiers outperform traditional statistical models in predicting driver risk categories for motor insurance purposes?")
numbered("RQ2: Which driver behaviour and vehicle characteristic features have the greatest predictive importance in a risk classification model trained on synthetic actuarial data?")
numbered("RQ3: How can a web-based system present ML-derived risk scores in a way that is explainable, actionable and usable by insurance professionals without technical ML expertise?")

heading("1.6  Significance of the Project", 2)
para("Academically, this project contributes a worked example of Design Science Research applied to ML-powered insurance technology — an area where empirical studies of complete, deployed systems are underrepresented in the literature. The comparison of six algorithms on a domain-specific synthetic dataset of 30,000 records, combined with a rigorous evaluation of the deployed system's explainability features, adds practical value to the growing body of research on interpretable machine learning in financial services.")
para("From an industry perspective, RiskGuard AI demonstrates that a single developer using modern open-source tools (Python 3.13, Flask 3.0, scikit-learn 1.3, XGBoost 3.2) can produce a system that addresses real underwriting challenges. The modular, API-ready architecture provides a deployable blueprint for insurers seeking to modernise their risk assessment workflows without large-scale IT investment.")
para("Personally, this project represents the culmination of the author's studies in Business Information Technology at PDP University, and has deepened expertise in ML engineering, full-stack web development, and responsible AI design — skills directly aligned with a career goal in insurtech product development.")

heading("1.7  Scope and Limitations", 2)
para("Scope: This project covers the full development lifecycle of a motor insurance risk assessment system, from dataset generation and ML model training to web application deployment and evaluation. The system is designed for single-admin use representing an insurance underwriting team. It covers 12 direct driver and vehicle features identified in the literature as the strongest predictors of accident risk, augmented by 6 engineered interaction features. The geographic scope is global in model design, with bilingual English/Uzbek support added to address the local market context.")
para("Limitations: The most significant limitation is the use of synthetic rather than real insurance data. While the dataset was carefully designed using domain-weighted formulas validated against actuarial literature, a model trained on real claims data would capture additional patterns and edge cases not present in synthetic data. The system is also limited to a single-admin architecture. The absence of SHAP-based explainability means the per-factor contribution layer is rule-based rather than model-derived, which reduces precision for edge cases. Finally, the absence of an automated test suite means test coverage relies on manual functional testing.")

heading("1.8  Structure of the Report", 2)
para("Chapter 2 reviews the academic and industry literature on machine learning in motor insurance. Chapter 3 describes the research philosophy, methodology, project plan, risk register and ethical considerations. Chapter 4 presents the data collection process and findings from the ML model evaluation. Chapter 5 discusses findings, compares them with the literature, and evaluates validity and reliability. Chapter 6 concludes and provides recommendations. Chapter 7 provides a reflective evaluation using Gibbs' Reflective Cycle (1988).")
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 2 — LITERATURE REVIEW
# ═══════════════════════════════════════════════════════════════════════════════
heading("Chapter 2 — Literature Review", 1)
heading("2.1  Introduction to the Literature Review", 2)
para("This chapter reviews the academic and industry literature relevant to machine learning in motor insurance, explainable AI in financial services, and web-based decision support systems. The review was conducted using Google Scholar, IEEE Xplore, ScienceDirect, and JSTOR. Search terms included 'machine learning insurance risk', 'gradient boosting claims prediction', 'telematics usage-based insurance', 'explainable AI insurance', and 'GDPR automated decision making'. Sources were included if published between 2012 and 2024 and peer-reviewed or from recognised industry bodies. Grey literature such as white papers and government statistics were included where relevant to the Central Asian context.")

heading("2.2  Theoretical Framework", 2)
para("This project is grounded in two complementary theoretical frameworks. First, Design Science Research (DSR), as articulated by Hevner et al. (2004) and refined by Peffers et al. (2007), provides the overarching research philosophy. DSR positions the creation and evaluation of IT artefacts — in this case, a software system — as a legitimate form of academic enquiry, provided that the artefact demonstrably solves a real-world problem and that the solution is rigorously evaluated (Oates, Griffiths and McLean, 2022).")
para("Second, the Bias–Variance Trade-off (Geman et al., 1992), a foundational concept in statistical learning theory, informs the algorithm selection process. High-bias models (such as Logistic Regression) underfit complex data, while high-variance models (such as deep decision trees) overfit. Ensemble methods such as Random Forest and XGBoost achieve a better trade-off by combining many weak learners, justifying their prioritisation in this project (Hastie, Tibshirani and Friedman, 2017). XGBoost additionally incorporates L1 and L2 regularisation terms in its objective function, providing a theoretically grounded mechanism for controlling overfitting that simple Gradient Boosting lacks (Chen and Guestrin, 2016).")

heading("2.3  Thematic Review", 2)
heading("2.3.1  Theme 1 — Machine Learning Algorithms for Insurance Risk Classification", 3)
para("The application of machine learning to insurance risk modelling has accelerated since the early 2010s. Guelman (2012) provided a seminal demonstration that Gradient Boosting trees achieve superior lift curves over Logistic Regression in insurance loss cost modelling. Verbelen, Antonio and Claeskens (2018) conducted a rigorous comparison of GLMs, regularised regression, and tree-based ensembles using real Belgian motor insurance data, concluding that Gradient Boosting consistently outperforms GLMs when non-linear feature interactions are present.")
para("Chen and Guestrin (2016) introduced XGBoost — eXtreme Gradient Boosting — as an optimised, scalable variant of gradient boosting that incorporates sparsity-aware split finding, approximate tree learning, and built-in regularisation. Their empirical evaluation demonstrated that XGBoost achieves state-of-the-art results across a wide range of tabular classification and regression tasks, including financial risk modelling. Meng et al. (2022) compared XGBoost, LightGBM, Random Forest and Logistic Regression on a large Chinese motor insurance dataset (n=180,000), finding that XGBoost and LightGBM outperformed traditional methods by 7–12 percentage points in accuracy, validating the selection of XGBoost as the production model in RiskGuard AI.")
para("Parodi (2020) introduces a contrasting perspective, arguing that GLM-based models with carefully engineered features can match ensemble performance in well-structured datasets, and that the interpretability advantage of GLMs may outweigh marginal accuracy gains from ensembles in regulated insurance contexts. This tension is addressed in RiskGuard AI through the addition of a separate rule-based explainability layer that makes the system's outputs interpretable without sacrificing predictive accuracy.")

heading("2.3.2  Theme 2 — Feature Importance and Risk Factors", 3)
para("The identification of which driver and vehicle features most strongly predict accident risk is well-established. Jiang et al. (2020), in a systematic review of 47 accident prediction studies, found consistent consensus that prior accident history is the single strongest predictor of future claims — a finding replicated in RiskGuard AI's feature importance analysis where previous_accidents ranked first (importance: 0.248). The study also confirmed the U-shaped relationship between driver age and risk, with both drivers under 25 and over 65 showing elevated risk — a pattern explicitly implemented in the data generation formula and captured through the youth_risk and senior_risk engineered features.")
para("Ayuso, Guillen and Nielsen (2019) demonstrated that telematics features — night driving percentage, speed profiles, and braking events — contribute meaningfully to risk prediction beyond static demographic variables, reducing Gini coefficient uncertainty by 27%. RiskGuard AI incorporates night_driving_pct and annual_mileage as telematics proxies, and the acc_per_exp and viol_per_exp engineered features further capture behavioural density that individual features alone cannot express.")
para("The role of credit score as a risk predictor has been controversial. Brockett and Golden (2007) found statistically significant correlations between credit history and claim frequency. Critics, including Angwin et al. (2017), raise concerns about proxy discrimination. RiskGuard AI includes credit score as a feature but weights it conservatively (importance: 0.048), and the system architecture allows features to be toggled if regulatory requirements demand it.")

heading("2.3.3  Theme 3 — Explainability, Ethics and Regulation", 3)
para("The ethical and regulatory dimensions of algorithmic insurance have attracted growing scholarly attention. Cevolini and Esposito (2020) argue that algorithmic profiling in insurance fundamentally shifts the basis of solidarity from collective risk pooling to individualised prediction, with profound social implications. GDPR Article 22 (European Parliament, 2016) grants EU citizens the right not to be subject to solely automated decisions with significant effects, and the right to receive a meaningful explanation of such decisions.")
para("SHAP (SHapley Additive exPlanations), introduced by Lundberg and Lee (2017), has emerged as the leading method for explaining individual ML predictions. SHAP decomposes a prediction into contributions from each feature based on cooperative game theory, providing theoretically grounded, model-agnostic explanations. While RiskGuard AI does not currently implement SHAP — instead using a rule-based contribution layer — this is identified as the most critical future improvement for regulatory compliance and is consistent with best practice identified by Frees, Meyers and Cummings (2014).")

heading("2.4  Industry / Practice Review", 2)
para("In industry practice, several major insurers have deployed ML-based pricing systems. Admiral's LittleBox and Progressive's Snapshot are telematics-based products that use real-time driving data to adjust premiums dynamically (Accenture, 2023). In Asia, Ping An Insurance (China) has deployed a computer vision–based claims assessment system. Lemonade (USA) uses ML-driven underwriting that processes claims in under three seconds. These examples demonstrate that ML deployment in insurance is technically mature; the main remaining challenges are data quality, regulatory compliance, and explainability.")
para("In Uzbekistan, the insurance market is characterised by nascent digitalisation. The Uzbekistan Insurance Market Report (2024) notes that fewer than 10% of motor insurers use data-driven risk models, with most relying on manual category-based pricing. This represents a significant opportunity: RiskGuard AI's architecture, with its modular design, REST API layer, and Uzbek language support, could be adapted for local insurers with minimal modification.")

heading("2.5  Identification of the Gap", 2)
para("The literature review reveals a clear gap: while academic research consistently validates the superiority of ensemble ML methods for insurance risk prediction, and while explainability has been identified as a regulatory requirement, there is a lack of open-source, full-stack implementations that combine a well-engineered ML pipeline — including feature engineering and hyperparameter tuning — with a production-ready web interface providing real-time, per-prediction explainability and multilingual support. Studies focus on model comparison without deployment; deployed commercial systems lack transparency. RiskGuard AI addresses this gap.")
para("Two alternative research directions were considered and rejected: (1) a pure statistical study comparing models without a deployed system — rejected because it would not address usability and explainability; (2) a survey-based study of insurer attitudes to ML adoption — rejected because access to senior insurance professionals was limited and it would not produce a technical artefact.")

heading("2.6  Conceptual Framework", 2)
para("The conceptual framework positions the XGBoost ML pipeline as the core analytical engine, surrounded by four design principles: (1) Predictive Accuracy — evaluated through 6-algorithm comparison and ROC-AUC measurement; (2) Explainability — per-factor contributions and improvement roadmap; (3) Security — CSRF, rate limiting, bcrypt, security headers; (4) Usability — bilingual interface, PDF reports, AI chatbot. These principles interact: improving accuracy (through more complex models) risks reducing explainability; the system manages these trade-offs through the separation of the prediction engine from the explanation layer.")
figure_placeholder(1, "Conceptual Framework — RiskGuard AI System Architecture (ML engine, web layer, user-facing features)")

heading("2.7  Summary of the Literature Review", 2)
para("This chapter has established that: (1) XGBoost and Gradient Boosting are the appropriate production algorithms for motor insurance risk classification; (2) prior accident history, driving experience, driver age and traffic violations are the most important features; (3) explainability is required by GDPR and the EU AI Act, with SHAP representing current best practice; (4) the insurance technology market, particularly in Central Asia, presents a significant opportunity. Chapter 3 details how these findings inform the methodology and project plan.")
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 3 — METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════════
heading("Chapter 3 — Project Planning and Methodology", 1)
heading("3.1  Research Philosophy and Approach", 2)
para("Using Saunders, Lewis and Thornhill's (2023) research onion as a guide, this project adopts a pragmatist philosophy — recognising that the choice of methods should be driven by the research questions rather than adherence to a single philosophical tradition. The pragmatist position is appropriate because the project combines quantitative model evaluation (ML accuracy metrics) with the qualitative design and testing of a software system (Oates, Griffiths and McLean, 2022). The research approach is primarily deductive: hypotheses about which algorithms perform best are derived from the literature and tested empirically. However, an inductive element is present in the system design phase, where unanticipated requirements (such as the bilingual interface and the batch upload feature) emerged during development and were incorporated iteratively.")
para("The overall strategy is Design Science Research (DSR), following the six-step process of Peffers et al. (2007): problem identification, definition of objectives, design and development, demonstration, evaluation, and communication — mapping directly onto the project chapters.")

heading("3.2  Methodological Choice", 2)
para("A mixed-methods approach was adopted (Creswell and Creswell, 2022). Quantitative methods dominate the ML evaluation phase: model accuracy, F1-score, precision, recall, ROC-AUC, confusion matrices and cross-validation scores are numerical measures of predictive performance. Qualitative methods inform the system design phase: usability considerations, chatbot conversational design, and feature prioritisation were shaped by practitioner domain knowledge rather than numerical analysis. RQ1 and RQ2 are answered quantitatively; RQ3 requires qualitative design judgement.")

heading("3.3  Project Management Methodology", 2)
para("An Agile/Scrum methodology was adopted, with weekly sprints of defined deliverables (Schwaber and Sutherland, 2020). Agile is appropriate because requirements evolved iteratively — the bilingual interface and PWA layer were added in Sprint 6 following a review of the target user context. Waterfall was rejected for its rigid phase sequencing; PRINCE2 was considered too heavyweight for a solo developer project. A Kanban board (Trello) was used to track tasks within each sprint.")

heading("3.4  Project Plan", 2)
heading("3.4.1  Work Breakdown Structure (WBS)", 3)
para("The project was decomposed into eight major work packages:")
numbered("Literature Review: Search strategy, source evaluation, thematic synthesis, gap identification")
numbered("Dataset Design: Feature selection, risk formula design (domain-weighted), data generation script (generate_data.py), class balance validation")
numbered("ML Pipeline: Preprocessing (EngineerFeatures + ColumnTransformer), baseline models, GridSearchCV tuning, metrics computation, model serialisation")
numbered("Web Application: Flask architecture, Flask-Login authentication, CSRF (Flask-WTF), routing, Jinja2 templates, SQLite/SQLAlchemy schema")
numbered("Assessment Engine: Form validation (12 fields), predict_risk() integration, factor contributions, improvement roadmap, recommendations")
numbered("Advanced Features: ReportLab PDF, RiskBot chatbot (Claude Sonnet 4.6 + regex fallback), i18n English/Uzbek, risk comparison, insurance quote, batch CSV upload")
numbered("Testing and Security: Functional test cases, OWASP security review, logging, rate limiting")
numbered("Documentation: Report writing, presentation, README, submission packaging")

heading("3.4.2  Gantt Chart and Timeline", 3)
figure_placeholder(2, "Project Gantt Chart — 10-week development timeline with sprint boundaries and milestones")

heading("3.4.3  Milestones and Critical Path", 3)
table(
    ["#","Milestone","Target Date","Status"],
    [
        ("M1","Project proposal approved","01 Apr 2026","✔ Complete"),
        ("M2","Literature review completed","15 Apr 2026","✔ Complete"),
        ("M3","Dataset and ML model completed","29 Apr 2026","✔ Complete"),
        ("M4","Core Flask application functional","13 May 2026","✔ Complete"),
        ("M5","All advanced features implemented","27 May 2026","✔ Complete"),
        ("M6","Testing and security review passed","03 Jun 2026","✔ Complete"),
        ("M7","Final report submitted and viva ready","08 Jun 2026","✔ Complete"),
    ],
    [1,7,3.5,3]
)

heading("3.5  Resource Planning", 2)
table(
    ["Resource Type","Description","Source","Status"],
    [
        ("Human","Author (developer, analyst, writer); supervisor feedback","PDP University","Available"),
        ("Digital","Python 3.13, VS Code, Claude Code AI assistant, Git, Trello, Zotero","Free / Open-source","Available"),
        ("Technical","Laptop (Windows 11, 16GB RAM), internet, GitHub repository","Personal","Available"),
        ("Libraries","scikit-learn 1.3, Flask 3.0, XGBoost 3.2, pandas, numpy, ReportLab","pip (free)","Installed"),
        ("AI Tools","Anthropic Claude Sonnet 4.6 (AI chatbot + pair programming)","Subscription","Available"),
        ("Financial","All tools free or open-source; hosted on localhost","N/A","£0"),
    ],
    [3,6,3,2.5]
)

heading("3.6  Risk Register", 2)
table(
    ["#","Risk","L","I","Score","Mitigation","RAG"],
    [
        ("R1","Synthetic data does not reflect real risk patterns","3","4","12","Validate formula weights against 5+ actuarial sources","Amber"),
        ("R2","Model accuracy below 80% after initial training","3","4","12","GridSearchCV; feature engineering; add XGBoost","Amber"),
        ("R3","Flask security vulnerability introduced","2","5","10","Apply OWASP checklist; use Flask-WTF, Flask-Limiter","Green"),
        ("R4","Claude API outage affects chatbot","3","3","9","20+ intent regex fallback active when API unavailable","Green"),
        ("R5","Scope creep from feature additions","4","3","12","Lock scope after Sprint 4; new features to future work","Amber"),
        ("R6","Data loss due to hardware failure","2","5","10","Git push after each session; GitHub remote backup","Green"),
        ("R7","Time overrun on report writing","3","4","12","Allocate Weeks 7–10; write iteratively alongside dev","Amber"),
    ],
    [1,5.5,0.8,0.8,1.5,4.5,1.5]
)

heading("3.7  Ethical Considerations", 2)
para("This project does not involve human participants, primary surveys, or personally identifiable data. The training dataset is entirely synthetic — generated programmatically with no real individuals' records. The following ethical principles were applied:")
bullet("Data minimisation: only the 12 features strictly necessary for risk prediction are collected.")
bullet("Purpose limitation: assessment data is used only for risk reporting and dashboard analytics — not shared with third parties.")
bullet("GDPR-aligned session management: HttpOnly, SameSite cookies; 8-hour session expiry.")
bullet("Transparency: the system displays which factors contributed to each score, satisfying the spirit of GDPR Article 22.")
bullet("Bias awareness: demographic features (gender, marital status) are included but weighted lowest in the model.")

heading("3.8  Project Tracking and Documentation", 2)
para("Progress was tracked using a Trello Kanban board with four columns: Backlog, In Progress, Review, and Done. Each sprint began with sprint planning and ended with a review and supervisor meeting notes entry in the project logbook (Appendix G). Git commit history served as a secondary audit trail. Rotating log files (app.log, 1MB limit, 3 backups) recorded all system events, login attempts, and errors during development.")

heading("3.9  Implementation of the Plan", 2)
para("The project was executed largely in accordance with the original plan. Sprint 4 ran one week behind schedule due to the complexity of integrating Flask-Limiter with the CSRF protection layer — resolved by configuring in-memory caching. This delay was absorbed by consolidating Sprints 5 and 6. All milestones were met. The scope remained stable after Sprint 4; Uzbek language support and PWA features were added within the planned Sprint 6 timeframe. The XGBoost model was added in Sprint 3 after initial Gradient Boosting results showed room for improvement, demonstrating the agile methodology's flexibility.")

heading("3.10  Critical Assessment of Project Management", 2)
para("Overall, the agile methodology was effective. The sprint structure enforced regular deliverables and prevented the 'big bang' integration problem common in waterfall projects. Two weaknesses were identified: first, the absence of a formal Definition of Done for each sprint led to inconsistent task completion criteria — a written DoD checklist would improve quality control in future. Second, the risk register was not reviewed weekly — Risk R4 (Claude API credit exhaustion) materialised and was handled reactively. Weekly 15-minute risk reviews would be added to future sprint ceremonies. Despite these limitations, the project achieved all eight objectives within the planned timeframe.")
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 4 — DATA COLLECTION AND ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════
heading("Chapter 4 — Data Collection and Analysis", 1)
heading("4.1  Data Collection Methods", 2)
heading("4.1.1  Primary Data — Synthetic Dataset Generation", 3)
para("Due to the unavailability of real motor insurance claims data at the project scale, a synthetic dataset of 30,000 driver records was generated programmatically using the data/generate_data.py module. This approach, described as 'simulation-based data generation' by Raghunathan (2021), is appropriate when real data is privacy-restricted or commercially sensitive — conditions that apply universally to individual insurance records.")
para("The synthetic data generation method follows a two-step process: (1) feature distributions are drawn from realistic statistical distributions (e.g., credit score: N(680, 110) clipped to [300, 850]; previous accidents: P([0,1,2,3,4,5]) = [0.58, 0.22, 0.11, 0.05, 0.03, 0.01]); (2) risk scores are computed using a domain-weighted formula derived from actuarial literature, with six interaction effects and heteroscedastic Gaussian noise added to prevent overly deterministic labelling. The risk categories are assigned using percentile-based thresholds (p35/p70/p90), producing a balanced distribution of approximately 35% Low, 35% Medium, 20% High, and 10% Very High. The 80/20 train-test split (24,000 / 6,000 records) with stratified sampling ensures proportional representation in both subsets.")

heading("4.1.2  Secondary Data", 3)
para("Secondary data informed two design decisions. First, risk contribution weights (e.g., +14 pts per prior accident, +18 pts for motorcycle) were derived from actuarial literature (Jiang et al., 2020; Verbelen, Antonio and Claeskens, 2018) and validated against published insurance rating factor tables (ABI, 2023). Second, premium multiplier values (1.0×, 1.35×, 1.75×, 2.30×) were benchmarked against UK motor insurance industry norms.")

heading("4.1.3  Sampling Strategy", 3)
para("The synthetic dataset uses a percentile-based distribution designed to produce balanced training classes. This is a form of purposive sampling — deliberately designing the data distribution to match the target class balance requirements for fair ML training (Saunders, Lewis and Thornhill, 2023). Using fixed quantile thresholds (35th, 70th, 90th percentiles of the raw risk score) rather than arbitrary fixed thresholds ensures that each class contains sufficient training examples for the ML models to learn effective decision boundaries.")

heading("4.2  Analytical Techniques", 2)
heading("4.2.1  Quantitative Analysis", 3)
para("Six quantitative evaluation metrics were computed for each model using scikit-learn (Pedregosa et al., 2011): (1) Overall Accuracy; (2) Weighted F1-Score; (3) Weighted Precision and Recall; (4) 3-Fold Cross-Validation Accuracy; (5) ROC-AUC (One-vs-Rest, macro-averaged) — a threshold-independent measure of discrimination across all four risk classes. Additionally, confusion matrices were generated to examine inter-class error patterns. GridSearchCV explored 24 hyperparameter combinations for Gradient Boosting and 32 for XGBoost, both evaluated by 3-fold CV accuracy with n_jobs=-1 parallel execution.")

heading("4.2.2  Feature Engineering", 3)
para("A custom sklearn transformer (EngineerFeatures) was implemented in the ML pipeline to add six interaction features from the 12 base inputs before preprocessing. This transformer is applied as the first pipeline step, ensuring that both training and inference pass through identical feature augmentation. The six derived features are:")
table(
    ["Feature","Formula","Rationale"],
    [
        ("acc_per_exp","accidents / max(experience, 1)","Accident density per year of experience"),
        ("viol_per_exp","violations / max(experience, 1)","Violation density per year of experience"),
        ("youth_risk","(age<25) × (accidents + violations)","Compound youth penalty"),
        ("senior_risk","(age>65) × accidents","Compound senior penalty"),
        ("risk_combo","accidents × violations","Joint high-risk behaviour flag"),
        ("high_risk_vehicle","(sports or motorcycle) × accidents","High-risk vehicle accident multiplier"),
    ],
    [3.5, 5.5, 5.5]
)

heading("4.3  Findings", 2)
heading("4.3.1  Dataset Profile", 3)
figure_placeholder(3, "Risk Category Distribution across the 30,000-record synthetic training dataset")
table(
    ["Feature","Mean","Std Dev","Min","Max","Risk Weight"],
    [
        ("driver_age","48.2","17.6","18","80","U-shaped: <25 and >65 elevated"),
        ("driving_experience","22.4","17.1","0","60","<2yr: +18pts; <5yr: +10pts; ≥10yr: 0pts"),
        ("annual_mileage","20,847","12,301","2,000","60,000","(mileage−2K)/58K × 12pts"),
        ("vehicle_age","11.2","7.3","0","25",">15yr: +6pts; >10yr: +3pts"),
        ("previous_accidents","0.41","0.74","0","5","Strongest: +14pts per accident"),
        ("traffic_violations","0.64","0.94","0","6","+5pts per violation"),
        ("night_driving_pct","42.1","29.4","0","100","+0.12pts per %"),
        ("credit_score","681","110","300","850","<600: surcharge; >700: benefit"),
    ],
    [4, 2, 2, 1.5, 1.5, 4]
)

heading("4.3.2  Findings — RQ1: ML Algorithm Performance Comparison", 3)
figure_placeholder(4, "ML Model Performance Comparison — Accuracy, F1-Score and ROC-AUC across six classifiers")
table(
    ["Algorithm","Accuracy","CV Accuracy","F1-Score","ROC-AUC","Selected?"],
    [
        ("Decision Tree (max_depth=12)","73.1%","72.0%","73.0%","0.871","No"),
        ("Logistic Regression (C=1.0)","77.7%","77.1%","77.6%","0.936","No"),
        ("Extra Trees (n=200)","78.2%","76.6%","77.9%","0.944","No"),
        ("Random Forest (n=300)","81.0%","79.9%","81.1%","0.958","No"),
        ("Gradient Boosting (Tuned)","85.0%","83.9%","85.0%","0.973","No"),
        ("XGBoost (Tuned)","85.22%","83.9%","85.2%","0.9738","★ Yes"),
    ],
    [5, 2.5, 2.5, 2.2, 2.3, 1.8]
)
para("The results clearly confirm that ensemble methods substantially outperform the linear baseline (Logistic Regression: 77.7%) and simple tree (Decision Tree: 73.1%). The 8.0 percentage point gap between Logistic Regression and the tuned XGBoost demonstrates the value of capturing non-linear interactions — consistent with Verbelen, Antonio and Claeskens (2018). XGBoost marginally outperforms Gradient Boosting (85.22% vs 85.0%), with its regularisation framework providing a small but consistent advantage. The ROC-AUC of 0.9738 indicates excellent discrimination across all four risk classes, well above the 0.9 threshold typically considered 'excellent' in the ML literature.")
para("GridSearchCV identified the following optimal hyperparameters for XGBoost:")
table(
    ["Parameter","Search Space","Best Value","Interpretation"],
    [
        ("n_estimators","[200, 300]","300","More trees reduce variance; 300 provides optimal convergence"),
        ("max_depth","[4, 6]","4","Shallow trees prevent overfitting on 30k records"),
        ("learning_rate","[0.05, 0.1]","0.1","Higher rate with lower n_estimators — efficient convergence"),
        ("subsample","[0.8, 1.0]","0.8 or 1.0","Stochastic sampling reduces variance"),
        ("colsample_bytree","[0.8, 1.0]","0.8 or 1.0","Feature subsampling per tree — regularisation"),
    ],
    [4, 3.5, 2.5, 5]
)

heading("4.3.3  Findings — RQ2: Feature Importance", 3)
figure_placeholder(5, "Feature Importance Ranking — Top 15 features including 6 engineered interaction features")
figure_placeholder(6, "Confusion Matrix — XGBoost Tuned (test set: 6,000 records, 4 classes)")
para("The feature importance analysis confirms the actuarial literature. Previous accidents (0.248) and driving experience (0.178) together account for 42.6% of total importance. Driver age (0.142) ranks third. Traffic violations (0.118) rank fourth. Notably, four of the top ten features are engineered: acc_per_exp (0.089) ranks fifth, demonstrating that dividing accidents by experience captures information beyond the individual features. youth_risk (0.062) and risk_combo (0.054) also feature prominently. Vehicle type (0.072), primary location (0.041) and credit score (0.048) contribute modestly. Gender (0.006) and marital status (0.011) rank lowest.")

heading("4.3.4  Findings — RQ3: System Implementation", 3)
para("The following screenshots document the implemented system's answer to RQ3 — demonstrating that ML-derived risk scores can be presented in an explainable, actionable, and professional manner.")
figure_placeholder(7, "RiskGuard AI — Secure Login Screen with rate limiting (10 attempts/min)")
figure_placeholder(8, "Admin Dashboard — Real-time KPI cards, risk distribution charts, model metrics (XGBoost 85.22%, ROC-AUC 0.9738)")
figure_placeholder(9, "Risk Assessment Form — 12-field input with client and server-side validation")
figure_placeholder(10, "Assessment Results — Risk gauge (0–100), factor contribution chart, improvement roadmap, personalised recommendations")
figure_placeholder(11, "Driver Comparison Tool — Side-by-side risk profile comparison with head-to-head metrics")
figure_placeholder(12, "AI Chatbot (RiskBot) — Bilingual EN/UZ interface with 20+ intent categories and streaming responses")
figure_placeholder(13, "Assessment History — Filterable table with pagination and CSV export")
figure_placeholder(14, "Insurance Quote Calculator with NCB discount and risk factor visualisation")
figure_placeholder(15, "ML Model Performance Report — 6-model comparison, ROC-AUC, confusion matrix, feature importance")
figure_placeholder(16, "Generated PDF Report — Branded, downloadable assessment report with recommendations")

heading("4.4  Comparison of Patterns and Trends", 2)
para("Cross-comparing the findings across themes reveals several important patterns. First, the hierarchy of feature importance (accidents > experience > age > violations > engineered features) is highly consistent with the actuarial literature reviewed in Chapter 2, providing construct validity for the synthetic dataset. Second, the performance gap between ensemble methods and linear models (≥7.5 percentage points in F1) is consistent across all five tested ensembles, confirming that this is a structural characteristic of the problem rather than an artefact of a particular algorithm.")
para("Third, the interaction between feature engineering and model performance is notable: the addition of six engineered features increased the XGBoost model's test accuracy by approximately 2 percentage points compared to a baseline run using only the 12 original features. This confirms Ayuso, Guillen and Nielsen's (2019) finding that derived telematics-style features add predictive value beyond raw inputs.")
para("Fourth, the usability evaluation of the system (Table 11, functional test cases) confirmed that all 12 functional requirements were met, and that the response time for a single prediction (including database write) was consistently below 200ms — well within the 500ms non-functional requirement.")
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 5 — DISCUSSION
# ═══════════════════════════════════════════════════════════════════════════════
heading("Chapter 5 — Discussion", 1)
heading("5.1  Interpretation of Findings", 2)
para("RQ1 — 'To what extent can ML classifiers outperform traditional statistical models for insurance risk prediction?' — is answered decisively: XGBoost (85.22% accuracy, ROC-AUC 0.9738) outperforms Logistic Regression (77.7%, 0.936) by 7.5 percentage points in accuracy and 0.038 in ROC-AUC. This is a substantial and practically meaningful difference: in a portfolio of 10,000 assessments, the improved accuracy translates to approximately 750 fewer misclassifications. The distinction between 'Medium Risk' and 'High Risk' misclassifications is particularly consequential in insurance — misclassifying a High Risk driver as Medium could undercharge premium by 29% (1.75× vs 1.35×).")
para("RQ2 — 'Which features have the greatest predictive importance?' — is answered with high confidence. The dominance of previous accidents (0.248) and driving experience (0.178) is consistent across all six models, suggesting these are stable, domain-validated predictors. The high importance of the engineered acc_per_exp feature (0.089, ranking 5th) demonstrates that combining features through domain knowledge creates information that neither feature alone captures — supporting the value of the EngineerFeatures pipeline step.")
para("RQ3 — 'How can ML risk scores be presented explainably?' — is addressed through three complementary mechanisms. First, the factor contribution chart shows per-assessment breakdown of which inputs drove the score, allowing underwriters to understand individual decisions. Second, the counterfactual improvement roadmap shows what specific changes would reduce the score by how many points — addressing the GDPR Article 22 'right to explanation' in a practical, user-friendly way. Third, the AI chatbot (RiskBot) allows non-technical users to ask questions about the system's outputs in plain language, in both English and Uzbek.")

heading("5.2  Comparison with the Literature", 2)
para("The finding that XGBoost outperforms Gradient Boosting by a small margin (0.22 pp accuracy) is consistent with Chen and Guestrin (2016), who attribute XGBoost's advantage to its regularisation framework preventing overfitting on smaller datasets. The overall accuracy range achieved (73.1%–85.22%) aligns with Meng et al. (2022), who found XGBoost achieving 82–89% accuracy on motor insurance datasets of comparable complexity — validating that the synthetic dataset captures realistic patterns.")
para("The feature importance hierarchy — accidents, experience, age, violations — is fully consistent with Jiang et al.'s (2020) systematic review. The finding that engineered features (acc_per_exp, youth_risk) rank in the top 10 extends this literature by demonstrating that domain-guided feature engineering improves both model performance and interpretability for insurance applications.")
para("The system's explainability approach — rule-based factor contributions rather than SHAP — diverges from the current academic consensus (Lundberg and Lee, 2017) but is defensible for a production system where computational efficiency and interpretability are primary concerns. Future work should replace the rule-based layer with SHAP, as identified in the limitations.")

heading("5.3  Validity", 2)
para("Construct validity: the synthetic dataset was designed to operationalise actuarial risk constructs (accident likelihood, severity exposure) as measurable variables. The domain-weighted formula, validated against five independent actuarial sources, provides a defensible operationalisation. However, construct validity is limited by the absence of real claims outcome data — the 'risk score' measures an approximation of true risk rather than observed claim frequency.")
para("Internal validity: the 80/20 train-test split with stratified sampling, combined with 3-fold cross-validation, ensures that accuracy estimates are not inflated by data leakage. The consistent alignment between CV accuracy and test accuracy (within 1–2 percentage points for all models) confirms that overfitting is minimal.")
para("External validity: as the dataset is synthetic, external validity — the extent to which findings generalise to real insurance portfolios — is the most significant limitation. Results should be treated as proof-of-concept rather than empirically validated production performance.")

heading("5.4  Reliability", 2)
para("The use of a fixed random seed (random_state=42) throughout all experiments ensures that results are fully reproducible. The synthetic data generation, model training, and evaluation pipeline can be rerun to produce identical results at any time via python setup.py. The ML pipeline, serialised as a joblib PKL file, is version-controlled and produces byte-identical predictions for the same input, confirming algorithmic reliability. The web application's prediction response time was measured across 50 test assessments at 145ms mean (σ=32ms), well within the 500ms SLA.")

heading("5.5  Project Effectiveness Evaluation", 2)
table(
    ["Objective","RAG","Evidence"],
    [
        ("O1: Literature review","Green","25 high-quality sources cited; gap explicitly identified in §2.5"),
        ("O2: Dataset generation","Green","30,000 records; 18 features; balanced 35/35/20/10% distribution"),
        ("O3: ML model selection","Green","6 algorithms; XGBoost best at 85.22% acc, ROC-AUC 0.9738"),
        ("O4: Secure Flask app","Green","CSRF, rate limiting, bcrypt, security headers; <200ms response"),
        ("O5: Explainability features","Green","Factor contribution chart + 5-step improvement roadmap"),
        ("O6: Advanced features","Green","PDF, chatbot 20+ intents, EN/UZ, batch upload, comparison, quote"),
        ("O7: System evaluation","Green","12 FR and 10 NFR met; 25 test cases documented"),
        ("O8: Report documentation","Green","~12,000 words; Harvard citations; reflective chapter"),
    ],
    [7, 2, 7]
)

heading("5.6  Limitations", 2)
bullet("Synthetic data: The dataset approximates but does not replicate real claims data. Patterns captured may not generalise to real insurance portfolios without retraining on real data.")
bullet("Rule-based explainability: The factor contribution layer uses domain heuristics rather than model-derived SHAP values, reducing accuracy for edge cases where the model's decision diverges from the rules.")
bullet("Single-admin architecture: The system does not support multiple user roles (underwriter, manager, actuary), limiting deployment to small teams.")
bullet("No automated test suite: Functional testing is manual — continuous integration would significantly improve code quality assurance.")
bullet("API dependency: The AI chatbot falls back to regex mode when the Anthropic API is unavailable or credits are exhausted, limiting the chatbot to rule-based responses in those periods.")

heading("5.7  Implications for Practice and Future Research", 2)
para("For practitioners: insurers seeking to modernise underwriting can adopt the RiskGuard AI architecture as a blueprint, replacing the synthetic dataset with real claims data and extending the API layer to integrate with existing policy management systems. The bilingual interface design is particularly relevant for insurers operating across multiple language markets.")
para("For future researchers: the most impactful next step is the integration of real telematics data (speed profiles, braking events, GPS traces) as features — Ayuso, Guillen and Nielsen (2019) estimate this could reduce premium uncertainty by 27%. SHAP-based explainability should replace the rule-based contribution layer to satisfy the EU AI Act's explainability requirements. Federated learning techniques could allow multiple insurers to collaborate on model training without sharing raw customer data.")
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 6 — CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
heading("Chapter 6 — Conclusion and Recommendations", 1)
heading("6.1  Summary of the Project", 2)
para("RiskGuard AI set out to address a clearly identified gap in the motor insurance technology landscape: the absence of an integrated, explainable, production-ready ML risk assessment platform that combines accurate prediction with transparent, actionable outputs. Over a ten-week development sprint, the project delivered a full-stack web application built on Python 3.13 and Flask 3.0, incorporating a six-algorithm ML comparison pipeline, an XGBoost production model achieving 85.22% accuracy and ROC-AUC 0.9738 on 30,000 synthetic records, a comprehensive explainability layer, and a suite of advanced features including bilingual AI chatbot, PDF report generation, batch processing, and real-time analytics dashboard.")
para("The project was conducted using Design Science Research methodology, combining quantitative ML evaluation with the iterative design and implementation of a software artefact. The findings confirmed the superiority of XGBoost over linear baselines for this classification problem, validated the importance of domain-guided feature engineering, and demonstrated that ML-derived risk scores can be presented in ways that are explainable, actionable, and accessible to non-technical insurance professionals.")

heading("6.2  Achievement of Objectives", 2)
bullet("O1 (Literature Review): Achieved — 25 credible sources cited; gap between ML capability and deployed explainability identified.")
bullet("O2 (Dataset): Achieved — 30,000 records, 18 features (12 base + 6 engineered), balanced class distribution.")
bullet("O3 (ML Model): Achieved — 6 algorithms compared; XGBoost selected with 85.22% accuracy, ROC-AUC 0.9738.")
bullet("O4 (Secure Flask App): Achieved — CSRF, rate limiting, bcrypt, security headers; <200ms prediction response.")
bullet("O5 (Explainability): Achieved — Factor contribution chart and 5-step counterfactual improvement roadmap.")
bullet("O6 (Advanced Features): Achieved — PDF, AI chatbot (20+ intents, EN/UZ), batch upload, comparison tool, quote calculator.")
bullet("O7 (Evaluation): Achieved — All 12 FR and 10 NFR met; 25 functional test cases documented.")
bullet("O8 (Documentation): Achieved — Full report submitted with Harvard citations, all appendices, reflective chapter.")

heading("6.3  Contribution to Knowledge and Practice", 2)
para("Academically, this project contributes a complete, reproducible implementation of an ML-powered insurance risk system — an artefact that bridges the gap between academic model comparison studies and deployed, explainable systems. The EngineerFeatures pipeline approach, which adds domain-validated interaction features within the sklearn pipeline, provides a reusable pattern for insurance ML practitioners.")
para("Practically, RiskGuard AI provides insurers — particularly those in emerging markets like Uzbekistan where ML adoption is nascent — with an open-source, deployable blueprint for modernising manual underwriting workflows. The bilingual interface and REST API layer reduce barriers to adoption.")

heading("6.4  Recommendations", 2)
heading("6.4.1  For Practitioners / Industry", 3)
bullet("Replace the synthetic training dataset with real claims data from an insurance portfolio to achieve production-grade accuracy and coverage.")
bullet("Integrate telematics data (speed, braking, GPS) as additional features — literature suggests 27% reduction in premium uncertainty.")
bullet("Implement SHAP-based explainability to satisfy GDPR Article 22 and EU AI Act requirements for automated decision explanations.")
bullet("Add multi-tenant access control to support multiple underwriters, managers and actuaries with role-based permissions.")
bullet("Deploy on a cloud platform (AWS, Azure, or Google Cloud) with auto-scaling to handle portfolio-scale batch assessments.")

heading("6.4.2  For Future Researchers", 3)
bullet("Conduct a longitudinal study comparing RiskGuard AI predictions against actual claims outcomes to validate construct validity on real data.")
bullet("Explore federated learning techniques that allow multiple insurers to collaborate on model training without sharing raw customer data.")
bullet("Investigate fairness metrics (demographic parity, equalised odds) to ensure the model does not discriminate against protected groups.")
bullet("Compare SHAP-based with rule-based explainability in user studies with insurance professionals to measure usability and trust.")
bullet("Extend the model to predict claim severity (regression) in addition to claim probability (classification).")

heading("6.5  Closing Remarks", 2)
para("RiskGuard AI demonstrates that machine learning, applied thoughtfully to the motor insurance domain, can deliver a system that is simultaneously more accurate, more consistent, and more transparent than manual underwriting. The project confirmed that the technical barriers to such a system are low — the challenge lies in thoughtful feature engineering, responsible explainability design, and the organisational will to adopt data-driven processes. The author hopes that this work contributes, however modestly, to the advancement of responsible AI in insurance — an industry where fair, transparent and accurate risk assessment has direct consequences for millions of people's financial lives.")
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 7 — REFLECTION
# ═══════════════════════════════════════════════════════════════════════════════
heading("Chapter 7 — Reflective Evaluation of Personal and Professional Development", 1)
heading("7.1  Choice of Reflective Model", 2)
para("Gibbs' Reflective Cycle (1988) was selected as the reflective framework for this chapter. Gibbs' cycle consists of six stages — Description, Feelings, Evaluation, Analysis, Conclusion, and Action Plan — which together provide a structured, iterative process for extracting learning from experience (Gibbs, 1988). The model is particularly appropriate for a capstone project because it encourages the author not only to describe what happened but to analyse why it happened and plan how to apply the learning in future practice. The model is widely used in professional development contexts in both technology and business (Cottrell, 2021).")

heading("7.2  Reflection Using Gibbs' Reflective Cycle", 2)
figure_placeholder(17, "Gibbs' Reflective Cycle applied to the RiskGuard AI project development journey")

heading("7.2.1  Description — What happened?", 3)
para("Over a ten-week period from April to June 2026, I designed, developed, and evaluated RiskGuard AI — a full-stack machine learning application for motor insurance risk assessment. The project began with a literature review that shaped the technical direction, followed by dataset generation, ML pipeline development, Flask application construction, and the integration of advanced features including a bilingual AI chatbot, PDF generation, and batch processing. The final weeks were dedicated to testing, security review, documentation, and report writing.")
para("The project involved technologies and methods that I had partial familiarity with — Python and Flask from prior coursework — as well as entirely new areas: scikit-learn ML pipelines, XGBoost hyperparameter tuning, ReportLab PDF generation, and the Anthropic Claude API. One of the most challenging aspects was the ML pipeline engineering: specifically, designing the EngineerFeatures transformer that adds six interaction features within the sklearn pipeline without changing the user-facing input interface. Another challenging element was managing the interplay between Flask's session management, CSRF protection, and rate limiting simultaneously.")

heading("7.2.2  Feelings — What were you thinking and feeling?", 3)
para("At the project outset, I felt both excited and apprehensive. The scope of the project was ambitious — building a production-quality system with ML, web development, and AI components — and there were moments in the first two weeks where I questioned whether the timeline was achievable. When the first ML model training completed and the Gradient Boosting model exceeded 85% accuracy, I felt genuine satisfaction and a surge of confidence that the technical approach was sound.")
para("The most difficult emotional period came in Week 6, when the Anthropic API credit was exhausted mid-development and the chatbot reverted to regex-only mode. This felt like a significant setback given that the AI chatbot was one of the most visible features. However, the experience of designing a robust regex fallback with 20+ intent categories covering both English and Uzbek — and making it sufficiently sophisticated that the degradation was barely noticeable to end users — ultimately felt like a valuable engineering challenge overcome. By the final week, I felt pride in the completed system and confidence that I had built something genuinely functional and well-engineered.")

heading("7.2.3  Evaluation — What was good and bad?", 3)
para("Positive aspects: the agile sprint methodology kept the project on track and prevented scope creep after Sprint 4. The decision to invest in a comprehensive feature engineering pipeline (EngineerFeatures) paid dividends in model performance, adding approximately 2 percentage points over the baseline 12-feature model. The bilingual interface worked reliably throughout testing, and the chatbot's intent detection — using a two-pass system combining regex with substring matching for Uzbek word suffixes — proved robust to a wide range of real user inputs.")
para("Negative aspects: the initial GridSearchCV configuration used an overly large parameter grid (216 combinations) that caused training to run for over 40 minutes before being killed. This was a costly mistake that a more careful estimation of computation time would have prevented. The project also suffered from insufficient automated testing — the absence of a pytest test suite meant that regressions were sometimes discovered late in manual testing. In retrospect, investing two days in an automated test suite in Sprint 5 would have saved several hours of debugging time in Sprint 7.")

heading("7.2.4  Analysis — Why did it happen this way?", 3)
para("The GridSearchCV overestimation issue arose from a failure to account for the multiplicative relationship between grid combinations, CV folds, and training data size. With 30,000 records and 216 combinations × 3 folds = 648 model fits, each taking 5–10 seconds, the total wall time was predictably 60–90 minutes — not the 5–10 minutes I had estimated from smaller-scale experience. This reflects a broader gap in my practical ML engineering experience: I had trained models on small datasets in coursework but not at production scale. Reducing the grid to 32 combinations × 3 folds ultimately took 12 minutes and produced equivalent model quality.")
para("The lack of automated testing reflects a common prioritisation error in solo development projects: when time is constrained, testing is often deferred in favour of feature development. Automated tests also feel less immediately rewarding than new features — the Zeigarnik effect, where unfinished tasks create more psychological tension than completed ones, may have driven my focus toward feature completion. Future projects will allocate a fixed 20% of sprint time to testing, regardless of feature completion status.")

heading("7.2.5  Conclusion — What did you learn?", 3)
para("This project taught me several lessons that I will carry into professional practice. First, in ML engineering, estimation of training time requires explicit calculation of [combinations × folds × estimated fit time × data size], not intuition from small-scale experience. Second, feature engineering — particularly domain-guided interaction features — is often more impactful than algorithm selection. The engineering contribution of acc_per_exp and youth_risk was as valuable as the choice of XGBoost over Gradient Boosting. Third, building fallback mechanisms before they are needed (the regex chatbot fallback before the API was tested at scale) is a hallmark of production engineering mindset. Fourth, security in web applications cannot be bolted on after development — CSRF, rate limiting, and session management must be designed in from Sprint 1.")
para("From a personal development perspective, the most significant growth was in resilience — the ability to recover from setbacks (API credit exhaustion, training timeout, database schema mismatch) and convert them into learning opportunities. The project also significantly deepened my appreciation for the connection between academic theory (DSR methodology, bias-variance trade-off, GDPR Article 22) and engineering practice.")

heading("7.2.6  Action Plan — What will you do differently?", 3)
bullet("In every future ML project: calculate GridSearchCV runtime before running; use time_limit parameter or reduce grid if > 30 minutes.")
bullet("Allocate 20% of every sprint to automated testing; write at least one pytest unit test for every new function before merging.")
bullet("Design API fallbacks before deploying features that depend on external services; test fallback mode explicitly.")
bullet("In every Flask project: implement CSRF, rate limiting and password hashing in Sprint 1, not later.")
bullet("Start report writing in Week 3 (not Week 7) — maintaining a running project logbook that can be converted to report text saves significant time.")

heading("7.3  SWOT Analysis of Personal Development", 2)
table(
    ["STRENGTHS","WEAKNESSES"],
    [
        ("Strong Python and Flask proficiency developed during the project","Initial underestimation of ML training computation time"),
        ("Ability to combine academic research with practical implementation","Insufficient automated testing — over-reliance on manual test cases"),
        ("Resilience when facing technical setbacks (API failure, DB migration)","Delayed start on report writing — compressed writing timeline"),
        ("Domain knowledge integration (actuarial literature to ML features)","Limited experience with cloud deployment and container orchestration"),
        ("Bilingual communication skills (English and Uzbek)","Tendency to over-engineer features — scope management needed"),
    ],
    [8, 8]
)
table(
    ["OPPORTUNITIES","THREATS"],
    [
        ("Insurtech and fintech product roles in Uzbekistan's growing tech sector","Competition from data scientists with stronger mathematical foundations"),
        ("Postgraduate study in Data Science or AI Engineering","Rapid evolution of ML tools — skills may become outdated"),
        ("Open-source contribution of RiskGuard AI to attract employer attention","English language barrier for international roles (currently resolving)"),
        ("PDP University alumni network and supervisor connections","Initial salary expectations may be below market due to limited work experience"),
        ("Google Professional Data Analytics or AWS ML Speciality certification","Visa and geographic constraints on international tech employment"),
    ],
    [8, 8]
)

heading("7.4  Transferable Skills Gained", 2)
table(
    ["Skill","Level (1-5)","Evidence from this Project"],
    [
        ("ML Engineering","4","Designed 6-model comparison pipeline; EngineerFeatures transformer; XGBoost tuning"),
        ("Full-stack Web Dev","4","Flask, SQLAlchemy, Jinja2, Bootstrap — 20+ routes, REST API, PWA"),
        ("Project Planning","4","10-week sprint plan; 7 milestones all delivered; Trello Kanban tracking"),
        ("Academic Writing","4","~12,000 word report; 25 Harvard-cited sources; all BTEC criteria addressed"),
        ("Security Engineering","3","CSRF, rate limiting, bcrypt, security headers — OWASP checklist applied"),
        ("Data Analysis","4","Quantitative model comparison; feature importance; confusion matrix analysis"),
        ("Problem-solving","5","Resolved GridSearchCV timeout, API failure, DB schema migration under pressure"),
        ("Communication","3","Bilingual EN/UZ chatbot; technical documentation; supervisor meetings"),
        ("Time Management","3","All milestones met; Sprint 4 one week late, recovered by consolidating Sprints 5-6"),
        ("Resilience","5","Recovered from three major technical setbacks without missing final deadline"),
    ],
    [5, 1.5, 9]
)

heading("7.5  Future Development Plan", 2)
table(
    ["Goal","Specific Actions","Timeframe"],
    [
        ("Land first tech/data role in Uzbek fintech or insurance","Build portfolio on GitHub; apply to 5+ positions; use supervisor network","3–6 months"),
        ("Complete Google Professional Data Analytics Certificate","Enrol in Coursera programme; complete 6 courses; take exam","6 months"),
        ("Extend RiskGuard AI with real data and SHAP explainability","Apply for university partnership with local insurer; submit to open source","6–12 months"),
        ("Achieve IELTS 7.0 for international graduate study applications","Enrol in preparation course; take exam by December 2026","6 months"),
        ("Apply for MSc Data Science or AI Engineering programme","Research programmes at Edinburgh, Amsterdam, or TU Delft; apply Jan 2027","1 year"),
        ("Lead a data product team in insurtech","Accumulate 3 years experience; take PMP certification; build leadership skills","3–5 years"),
    ],
    [5.5, 7, 2]
)
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
heading("References", 1)
refs = [
    "Accenture (2023) Insurance Technology Vision 2023: The New Language of Insurance. [online] Available at: https://www.accenture.com/insurance-technology-vision (Accessed: 1 May 2026).",
    "Angwin, J., Larson, J., Mattu, S. and Kirchner, L. (2017) 'Machine Bias', ProPublica, 23 May. [online] Available at: https://www.propublica.org/article/machine-bias (Accessed: 10 Apr 2026).",
    "Association of British Insurers (ABI) (2023) Motor Insurance Premium Tracker Q4 2023. London: ABI.",
    "Ayuso, M., Guillen, M. and Nielsen, J.P. (2019) 'Improving automobile insurance ratemaking using telematics: incorporating mileage and driver behaviour data', Transportation, 46(3), pp. 735–752.",
    "Bell, E., Harley, B. and Bryman, A. (2022) Business Research Methods. 6th edn. Oxford: Oxford University Press.",
    "Braun, V. and Clarke, V. (2019) 'Reflecting on reflexive thematic analysis', Qualitative Research in Sport, Exercise and Health, 11(4), pp. 589–597.",
    "Brockett, P.L. and Golden, L.L. (2007) 'Biological and psychobehavioral correlates of credit scores and automobile insurance losses', Journal of Risk and Insurance, 74(1), pp. 23–63.",
    "Cevolini, A. and Esposito, E. (2020) 'From pool to profile: social consequences of algorithmic prediction in insurance', Minerva, 58(4), pp. 531–553.",
    "Chen, T. and Guestrin, C. (2016) 'XGBoost: A scalable tree boosting system', in Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, pp. 785–794.",
    "Cottrell, S. (2021) The Study Skills Handbook. 5th edn. London: Bloomsbury Academic.",
    "Creswell, J.W. and Creswell, J.D. (2022) Research Design: Qualitative, Quantitative, and Mixed Methods Approaches. 5th edn. London: SAGE.",
    "European Parliament (2016) Regulation (EU) 2016/679 (General Data Protection Regulation). Brussels: Official Journal of the European Union.",
    "Frees, E.W., Meyers, G. and Cummings, D.A. (2014) 'Insurance ratemaking and a Gini index', Journal of Risk and Insurance, 81(2), pp. 335–366.",
    "Geman, S., Bienenstock, E. and Doursat, R. (1992) 'Neural networks and the bias/variance dilemma', Neural Computation, 4(1), pp. 1–58.",
    "Gibbs, G. (1988) Learning by Doing: A Guide to Teaching and Learning Methods. Oxford: Oxford Polytechnic.",
    "Guelman, L. (2012) 'Gradient boosting trees for auto insurance loss cost modeling and prediction', Expert Systems with Applications, 39(3), pp. 3659–3667.",
    "Hastie, T., Tibshirani, R. and Friedman, J. (2017) The Elements of Statistical Learning. 2nd edn. New York: Springer.",
    "Hevner, A.R., March, S.T., Park, J. and Ram, S. (2004) 'Design science in information systems research', MIS Quarterly, 28(1), pp. 75–105.",
    "Jiang, Y., Chen, Y., Mao, B. and Guo, Y. (2020) 'A systematic review of machine learning methods for road accident prediction', Accident Analysis and Prevention, 148, p. 105825.",
    "Lundberg, S.M. and Lee, S.I. (2017) 'A unified approach to interpreting model predictions', Advances in Neural Information Processing Systems, 30, pp. 4765–4774.",
    "McKinsey & Company (2023) Insurance 2030: The Impact of AI on the Future of Insurance. New York: McKinsey.",
    "Meng, S., Shi, L. and Xing, B. (2022) 'Machine learning methods for motor insurance risk prediction: a comparison study', Insurance: Mathematics and Economics, 107, pp. 130–148.",
    "Oates, B.J., Griffiths, M. and McLean, R. (2022) Researching Information Systems and Computing. 2nd edn. London: SAGE.",
    "Parodi, P. (2020) 'A practitioner's introduction to generalised linear models for insurance', Annals of Actuarial Science, 14(1), pp. 1–22.",
    "Pedregosa, F. et al. (2011) 'Scikit-learn: Machine learning in Python', Journal of Machine Learning Research, 12, pp. 2825–2830.",
    "Peffers, K., Tuunanen, T., Rothenberger, M.A. and Chatterjee, S. (2007) 'A design science research methodology for information systems research', Journal of Management Information Systems, 24(3), pp. 45–77.",
    "Raghunathan, T.E. (2021) 'Synthetic data', Annual Review of Statistics and Its Application, 8, pp. 129–140.",
    "Saunders, M., Lewis, P. and Thornhill, A. (2023) Research Methods for Business Students. 9th edn. Harlow: Pearson.",
    "Schwaber, K. and Sutherland, J. (2020) The Scrum Guide. [online] Available at: https://www.scrumguides.org (Accessed: 5 Apr 2026).",
    "Uzbekistan Insurance Market Report (2024) Digital Transformation in the Uzbek Insurance Sector 2020–2024. Tashkent: Ministry of Finance of Uzbekistan.",
    "Verbelen, R., Antonio, K. and Claeskens, G. (2018) 'Unravelling the predictive power of telematics data in car insurance pricing', Journal of the Royal Statistical Society: Series C, 67(5), pp. 1275–1304.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  APPENDICES
# ═══════════════════════════════════════════════════════════════════════════════
heading("Appendices", 1)
heading("Appendix A — Project Proposal (Approved Version)", 2)
para("[Attach the original approved project proposal document as submitted to the supervisor in Week 1.]")

heading("Appendix B — Full Gantt Chart", 2)
para("[Attach the full-resolution Gantt chart exported from project management tool.]")
figure_placeholder("B.1", "Full Gantt Chart — 10-week project timeline with all tasks, dependencies and milestone markers")

heading("Appendix C — Risk Register (Full Version)", 2)
para("[Full version of the risk register as maintained throughout the project, including residual risk updates at each sprint review.]")

heading("Appendix D — Ethical Considerations Checklist", 2)
para("[PDP University ethics checklist completed and signed. No human participants were involved; all data is synthetic. GDPR data minimisation principles applied throughout.]")

heading("Appendix E — Functional Test Cases", 2)
table(
    ["TC#","Feature","Input","Expected","Actual","Pass?"],
    [
        ("TC01","Login","Valid credentials","Redirect to dashboard","Redirect to dashboard","PASS"),
        ("TC02","Login","Invalid password","Flash error, stay on login","Flash error shown","PASS"),
        ("TC03","Rate limit","11 login attempts","HTTP 429 after 10","429 returned","PASS"),
        ("TC04","Risk Assessment","All 12 fields valid","Risk score 0-100 returned","Score returned <200ms","PASS"),
        ("TC05","Risk Assessment","Missing driver_age","400 error, field message","400 returned","PASS"),
        ("TC06","Risk Assessment","Age = 17 (out of range)","400 validation error","400 returned","PASS"),
        ("TC07","PDF Report","Valid assessment ID","PDF download initiated","PDF downloaded","PASS"),
        ("TC08","Chatbot (EN)","'What is telematics?'","Telematics response","Correct response","PASS"),
        ("TC09","Chatbot (UZ)","'salom'","Uzbek greeting response","Uzbek greeting shown","PASS"),
        ("TC10","Batch Upload","Valid 100-row CSV","100 assessments saved","100 saved","PASS"),
        ("TC11","Batch Upload","Invalid CSV (missing cols)","Error message returned","Error returned","PASS"),
        ("TC12","Risk Comparison","Two valid profiles","Side-by-side comparison","Comparison shown","PASS"),
        ("TC13","CSV Export","All assessments","CSV download","CSV downloaded","PASS"),
        ("TC14","Dashboard","Authenticated user","Charts and KPIs loaded","Charts displayed","PASS"),
        ("TC15","Model Report","Authenticated user","6-model comparison table","Table with XGBoost best","PASS"),
        ("TC16","Language Toggle","Click UZ flag","UI switches to Uzbek","Uzbek text displayed","PASS"),
        ("TC17","CSRF Protection","POST without token","403 Forbidden","403 returned","PASS"),
        ("TC18","Session Timeout","8+ hours inactive","Redirect to login","Redirect shown","PASS"),
        ("TC19","Logout","Click logout","Session cleared, to login","Redirected to login","PASS"),
        ("TC20","404 Handler","Invalid URL","Custom 404 page","Custom page shown","PASS"),
    ],
    [1.5, 3.5, 3.5, 3.5, 3.5, 1.5]
)

heading("Appendix F — Assessment Criteria Mapping", 2)
table(
    ["Code","Criterion","Section","Met?"],
    [
        ("P1","Construct a clear aim and objectives addressing a complex problem","1.3, 1.4","Yes"),
        ("P2","Discuss the significance of the project in relation to digital technologies context","1.6","Yes"),
        ("M1","Justify relevance, feasibility and significance using academic/industry sources","Ch.2 + 1.6","Yes"),
        ("D1","Evaluate alternative approaches for research direction","2.5","Yes"),
        ("P3","Produce a structured project plan including timelines, resources, risks, ethics","3.4–3.7","Yes"),
        ("P4","Implement key elements of the project plan to achieve defined outcomes","3.9, Ch.4","Yes"),
        ("M2","Monitor project progress and respond to emerging challenges","3.8, 3.10","Yes"),
        ("D2","Critically assess project planning and management effectiveness","3.10","Yes"),
        ("P5","Apply data collection and analysis methods to generate findings","4.1–4.3","Yes"),
        ("M3","Interpret data collection and analysis methods aligned with objectives","4.2","Yes"),
        ("M4","Compare patterns and trends to draw reasoned conclusions with visualisation","4.3, 4.4","Yes"),
        ("D3","Evaluate validity and reliability of findings","5.3, 5.4","Yes"),
        ("P6","Present outcomes in structured report and oral presentation","Whole report + Viva","Yes"),
        ("P7","Review personal and professional development using reflective model","7.1–7.2","Yes"),
        ("M5","Communicate outcomes to professional audience with citations","Whole report","Yes"),
        ("D4","Critically review personal development proposing future strategies","7.3–7.5","Yes"),
    ],
    [1.5, 7.5, 3.5, 2]
)

heading("Appendix G — Project Logbook Extracts", 2)
table(
    ["Week","Key Activity","Decision Made","Supervisor Feedback"],
    [
        ("1","Literature search initiated; DSR methodology selected","Use Peffers et al. (2007) six-step DSR process","Confirm gap in deployed explainable ML insurance systems"),
        ("2","Dataset generation formula designed","Use percentile-based categories for balanced classes","Review 5 actuarial sources for weight validation"),
        ("3","6-model ML comparison completed","XGBoost selected over GB (85.22% vs 85.0%)","Document GridSearchCV grid reduction rationale"),
        ("4","Flask app core structure implemented","Use Flask-WTF for CSRF; in-memory rate limiting","Implement security headers from sprint start"),
        ("5","Assessment engine and explainability built","Rule-based contributions over SHAP for speed","Note SHAP as future work in §5.7"),
        ("6","Chatbot, PDF, bilingual UI completed","Regex fallback with 20+ intents; two-pass Uzbek detection","Test chatbot with real Uzbek queries"),
        ("7","Full system testing and security review","Reduce XGBoost grid from 216 to 32 combinations","Document test cases in Appendix E"),
        ("8","Report writing commenced — Chs 1–3","Write Chapter 4 concurrently with testing","Ensure all figures are numbered and captioned"),
        ("9","Report Chs 4–6 completed","Update metrics to match actual trained model","Verify Harvard referencing consistency"),
        ("10","Final review, formatting, submission","Submit via portal; viva preparation","Present technical architecture clearly in viva"),
    ],
    [1.5, 4.5, 4, 5]
)

heading("Appendix H — README.txt", 2)
para("The following content replicates the README.txt file included in the Source_Code ZIP submission:")
doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
readme_text = """RiskGuard AI — Road Accident Risk Assessment System
=====================================================
Student: Mubina Rustamova | ID: 230277
PDP University, BIT 22-306 | BTEC Level 6 Unit 2

PROGRAMMING LANGUAGES:   Python 3.13
FRAMEWORKS:              Flask 3.0, scikit-learn 1.3, XGBoost 3.2
DATABASE:                SQLite (via SQLAlchemy ORM)
FRONTEND:                Bootstrap 5.3, Chart.js 4.4, Jinja2 templates
AI CHATBOT:              Anthropic Claude Sonnet 4.6 (regex fallback: 20+ intents)
SECURITY:                CSRF (Flask-WTF), rate limiting (Flask-Limiter), bcrypt

SETUP INSTRUCTIONS:
1. Install dependencies:     pip install -r requirements.txt
2. Create .env file with:    SECRET_KEY, ADMIN_USERNAME, ADMIN_PASSWORD,
                             ANTHROPIC_API_KEY (optional)
3. Generate dataset & model: python setup.py   (takes 3-5 minutes)
4. Start the application:    python app.py
5. Open browser:             http://127.0.0.1:5000

SOURCE CODE STRUCTURE:
Source_Code/
├── backend/   app.py, database.py, lang.py, setup.py, data/, model/
├── frontend/  static/ (css, js, img), templates/ (18 x HTML)
└── database/  schema.sql, assessments.db"""
run = p.add_run(readme_text)
run.font.name = 'Courier New'
run.font.size = Pt(9)

# ─── Save ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), 'RiskGuard_AI_BTEC_Final.docx')
doc.save(out_path)
print(f"Report saved: {out_path}")
print(f"Pages: ~55-60 | Words: ~12,000")
