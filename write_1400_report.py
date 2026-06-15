# -*- coding: utf-8 -*-
# 1400-WORD DEFENSE / VIVA REPORT GENERATOR — ENGLISH
# RiskGuard AI — Pearson BTEC Level 6 | PDP University
# Run:  python write_1400_report.py

import os, re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT   = r"C:/Users/rusta/Desktop/RiskGuard_AI_DEFENSE_EN.docx"
GITHUB   = "https://github.com/mubinarustamova344-netizen/riskguard-ai"
LIVE_URL = "https://riskguard-ai.onrender.com"

doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.0)

doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)

NAVY  = RGBColor(13,  27,  42)
GOLD  = RGBColor(245,158, 11)
BLUE  = RGBColor(37,  99, 235)
GRAY  = RGBColor(71,  85, 105)
WHITE = RGBColor(255,255,255)
GREEN = RGBColor(16, 163, 74)

def pb():
    doc.add_page_break()

def ctr(text, size=13, bold=False, italic=False, color=None):
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pp.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return pp

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    hd.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if hd.runs:
        hd.runs[0].font.name = 'Times New Roman'
    return hd

def p(text, bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=None):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.line_spacing = Pt(20)
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return para

def bullet(text, size=12):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Cm(1.0)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return para

def tbl_hdr(tbl, headers):
    row = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = row.cells[i]
        cell.text = hdr
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = WHITE
        try:
            tc  = cell._tc
            tcP = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  '0D1B2A')
            tcP.append(shd)
        except Exception:
            pass

def tbl_row(tbl, values):
    row = tbl.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'

def add_hyperlink(para, url, text=None):
    """Add a clickable hyperlink run to a paragraph."""
    REL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink'
    rel_id = para.part.relate_to(url, REL, is_external=True)
    hlink  = OxmlElement('w:hyperlink')
    hlink.set(qn('r:id'), rel_id)
    h_r   = OxmlElement('w:r')
    h_rPr = OxmlElement('w:rPr')
    c  = OxmlElement('w:color'); c.set(qn('w:val'), '0563C1'); h_rPr.append(c)
    u  = OxmlElement('w:u');     u.set(qn('w:val'), 'single'); h_rPr.append(u)
    sz = OxmlElement('w:sz');    sz.set(qn('w:val'), '24');    h_rPr.append(sz)
    h_r.append(h_rPr)
    t_el = OxmlElement('w:t')
    t_el.text = text or url
    h_r.append(t_el)
    hlink.append(h_r)
    para._p.append(hlink)
    return hlink

def link_line(label, url, display=None):
    """A labelled line with a clickable URL."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1.0)
    para.paragraph_format.space_after = Pt(4)
    r1 = para.add_run(f"{label}: ")
    r1.bold = True; r1.font.size = Pt(12); r1.font.name = 'Times New Roman'
    add_hyperlink(para, url, display or url)

# ═══════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════
for _ in range(2):
    doc.add_paragraph()

ctr("PDP UNIVERSITY", 15, bold=True, color=NAVY)
ctr("Faculty of Artificial Intelligence (AI) Solutions and Applications", 12, bold=True)
ctr("Tashkent, Uzbekistan", 11, italic=True, color=GRAY)
doc.add_paragraph()
ctr("Pearson BTEC Level 6 Diploma in Digital Technologies (SRF)", 11, italic=True, color=BLUE)
ctr("Unit 2 — Independent Project (70726U)", 11, italic=True)
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run(
    "RiskGuard AI: A Machine Learning–Powered Road Accident\n"
    "Risk Assessment System for Insurance Companies"
)
tr.bold = True
tr.font.name = 'Times New Roman'
tr.font.size = Pt(16)
tr.font.color.rgb = NAVY

doc.add_paragraph()
ctr("Defense / Viva Summary Report", 13, italic=True, color=GRAY)
doc.add_paragraph()

for label, val in [
    ("Student Name:",      "Mubina Rustamova"),
    ("Student ID:",        "230277"),
    ("Programme / Group:", "Business Information Technology — 22-306"),
    ("Supervisor:",        "Xolmirzayev Muxammadjon"),
    ("Submission Date:",   "13.06.2026"),
    ("Word Count:",        "~2,000 words"),
]:
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = pp.add_run(f"{label}  ")
    r1.bold = True; r1.font.size = Pt(12); r1.font.name = 'Times New Roman'
    r2 = pp.add_run(val)
    r2.font.size = Pt(12); r2.font.name = 'Times New Roman'

# Live demo and GitHub as hyperlinks on title page
for label, url in [("Live Demo:", LIVE_URL), ("GitHub:", GITHUB)]:
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = pp.add_run(f"{label}  ")
    r1.bold = True; r1.font.size = Pt(12); r1.font.name = 'Times New Roman'
    add_hyperlink(pp, url)

doc.add_paragraph()
ctr("Academic Year 2025–2026", 11, italic=True, color=GRAY)
pb()

# ═══════════════════════════════════════════════
# CHAPTER 1 — INTRODUCTION
# ═══════════════════════════════════════════════
h("1.  Introduction and Problem Statement")

p("Despite the explosion of machine learning tools over the past decade, most motor insurers — "
  "particularly in emerging markets such as Uzbekistan — still assess driver risk using narrow "
  "demographic variables: age, gender and postcode. The result is pricing that is inconsistent, "
  "opaque and frustrating for both insurer and policyholder. A 45-year-old accountant driving "
  "40,000 miles a year at night with two prior accidents carries far higher risk than one driving "
  "8,000 miles in daylight with a clean record — yet traditional models cannot distinguish them.")

p("The global context underscores the urgency. The ABI (2023) reported UK motor insurance claims "
  "of £6.7 billion in 2023, with £1.1 billion attributable to fraud. McKinsey and Company (2023) "
  "estimate that ML-powered underwriting can reduce premium uncertainty by 20–35% compared to "
  "traditional models. In Uzbekistan, the motor insurance sector grew at 18% CAGR (2020–2024) "
  "following regulatory liberalisation, yet fewer than 10% of local insurers use data-driven "
  "pricing models (Uzbekistan Insurance Market Report, 2024) — a gap this project directly addresses.")

p("The four interconnected problems this project solves are: (1) too-narrow feature space — fewer "
  "than ten demographic variables; (2) poor consistency — different underwriters make different "
  "decisions for identical profiles; (3) no actionable feedback for policyholders; and (4) lack of "
  "explainability required by GDPR Article 22 and the EU AI Act (2024).")

# ═══════════════════════════════════════════════
# CHAPTER 2 — AIM AND OBJECTIVES
# ═══════════════════════════════════════════════
h("2.  Project Aim and SMART Objectives")

p("The aim of this project is to design, develop and evaluate a machine-learning–powered web "
  "application — RiskGuard AI — that automates motor insurance driver risk profiling, generates "
  "explainable risk scores and premium multipliers, and provides personalised improvement "
  "recommendations for insurance underwriters.")

tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
tbl_hdr(tbl, ["Ref", "Objective", "Success Criterion", "Target"])
for row in [
    ("O1","Review ≥20 academic and industry sources on ML in motor insurance",
     "Gap identified; all sources critically evaluated","Week 2"),
    ("O2","Design and generate a realistic synthetic training dataset",
     "12,000 records, 12 features, <15% class imbalance","Week 3"),
    ("O3","Train, compare and select the best ML classifier",
     "≥4 algorithms; best >85% F1-score; GridSearchCV tuning","Week 4"),
    ("O4","Build a secure, full-stack Flask web application",
     "CSRF, rate limiting, PBKDF2 hashing; <500ms response","Week 7"),
    ("O5","Implement per-prediction explainability",
     "Factor contributions breakdown + top-5 improvement roadmap","Week 8"),
    ("O6","Deliver AI chatbot, PDF export and bilingual EN/UZ interface",
     "20+ intent categories; Claude Haiku integration","Week 9"),
    ("O7","Evaluate against functional and non-functional requirements",
     "All 12 FR + 10 NFR met; OWASP security review passed","Week 9"),
    ("O8","Produce a BTEC Level 6 academic report",
     "≥14,000 words; Harvard citations; all criteria addressed","Week 10"),
]:
    tbl_row(tbl, row)
doc.add_paragraph()

# ═══════════════════════════════════════════════
# CHAPTER 3 — LITERATURE REVIEW
# ═══════════════════════════════════════════════
h("3.  Literature Review and Theoretical Framework")

p("The review is structured around three themes corresponding to the research questions: "
  "(1) ML algorithm selection for insurance risk classification; (2) feature importance and "
  "actuarial evidence; (3) explainability, ethics and regulation. Sources were drawn from "
  "Google Scholar, IEEE Xplore and ScienceDirect (2012–2024).")

p("Guelman (2012) was among the first to apply Gradient Boosting to insurance loss cost "
  "modelling, demonstrating superior lift curves over Logistic Regression. Verbelen, Antonio "
  "and Claeskens (2018) analysed over 150,000 Belgian policies and found Gradient Boosting "
  "outperforms GLMs by 8–15 percentage points when non-linear feature interactions are present — "
  "the most directly relevant source for algorithm selection in this project.")

p("Jiang et al. (2020) reviewed 47 independent accident-prediction studies and established that "
  "prior accident history is the single strongest predictor of future claims across all methods "
  "and geographic contexts — directly motivating the heavy weight (+14 points per accident) given "
  "to previous_accidents in RiskGuard AI's risk formula.")

p("GDPR Article 22 (European Parliament, 2016) grants citizens the right to a meaningful "
  "explanation of automated decisions with significant effects. Lundberg and Lee (2017) proposed "
  "SHAP (SHapley Additive exPlanations), the theoretical gold standard for tree-model explanation. "
  "RiskGuard AI currently uses a rule-based contribution layer (simpler, more accessible to "
  "non-technical underwriters) but identifies SHAP implementation as the top future upgrade for "
  "regulatory compliance.")

# ═══════════════════════════════════════════════
# CHAPTER 4 — METHODOLOGY
# ═══════════════════════════════════════════════
h("4.  Research Methodology")

p("The research philosophy is pragmatism (Saunders, Lewis and Thornhill, 2023): methods are "
  "driven by research questions rather than philosophical orthodoxy. The overall strategy is "
  "Design Science Research (DSR, Peffers et al., 2007), which treats software system creation "
  "and evaluation as legitimate academic research — the correct choice because the research "
  "questions cannot be fully answered without a working deployed system. Project management "
  "follows Agile/Scrum with 10 weekly sprints.")

p("The training dataset is synthetic — 12,000 records, 12 features, generated using a "
  "domain-weighted actuarial formula validated against WHO statistics and Jiang et al. (2020). "
  "Real insurance claims data is commercially sensitive and inaccessible to individual researchers "
  "without institutional agreements. Synthetic data allows a full technical demonstration while "
  "keeping the pipeline reproducible (random seed = 42). The class distribution — Low 25.3%, "
  "Medium 30.1%, High 27.8%, Very High 16.8% — mirrors published industry portfolio statistics.")

# ═══════════════════════════════════════════════
# CHAPTER 5 — SYSTEM IMPLEMENTATION
# ═══════════════════════════════════════════════
h("5.  System Implementation")
h("5.1  ML Pipeline and Model Selection", 2)

p("The scikit-learn 1.4 Pipeline chains StandardScaler (numeric features) and OneHotEncoder "
  "(categorical features) with the classifier. Four algorithms were trained on the 9,600-record "
  "training set and evaluated on the 2,400-record held-out test set. GridSearchCV explored 108 "
  "hyperparameter combinations (3×3×3×2) using 5-fold cross-validation.")

tbl2 = doc.add_table(rows=1, cols=6)
tbl2.style = 'Table Grid'
tbl_hdr(tbl2, ["Algorithm", "Accuracy", "CV Accuracy", "F1-Score", "Precision", "Selected?"])
for row in [
    ("Decision Tree (max_depth=10)",  "84.2%", "82.1%", "83.8%", "84.0%", "No"),
    ("Logistic Regression (C=1.0)",   "79.1%", "78.4%", "78.6%", "78.9%", "No"),
    ("Random Forest (n=200)",          "88.7%", "87.3%", "88.4%", "88.6%", "No"),
    ("Gradient Boosting (Tuned) ★",   "91.3%", "90.1%", "91.0%", "91.2%", "YES ✓"),
]:
    tbl_row(tbl2, row)
doc.add_paragraph()
p("Table 1: ML model comparison results (test set: 2,400 records). All metrics are weighted averages.",
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
doc.add_paragraph()

p("Gradient Boosting (91.3%) outperforms Logistic Regression (79.1%) by 12.2 percentage points, "
  "confirming that the non-linear relationships in the data — the U-shaped driver age curve, the "
  "multiplicative interaction between accidents and experience, and vehicle-type threshold effects — "
  "cannot be captured by a linear model. The small test–CV accuracy gap (max 2.1% across all "
  "models) confirms no significant overfitting. The selected model configuration: n_estimators=150, "
  "max_depth=5, learning_rate=0.08, subsample=0.8.")

h("5.2  Security Architecture", 2)

p("Security was designed in from Sprint 3, not added as an afterthought. The six-layer architecture "
  "addresses the OWASP Top 10 (2021): Layer 1 — Authentication (PBKDF2-SHA256 password hashing, "
  "260,000 iterations; credentials in .env excluded from Git); Layer 2 — CSRF protection (Flask-WTF, "
  "cryptographic token on every POST); Layer 3 — Rate limiting (Flask-Limiter, 10 login attempts "
  "per minute per IP, HTTP 429 thereafter); Layer 4 — SQL injection prevention (SQLAlchemy ORM, "
  "parameterised queries only); Layer 5 — XSS prevention (Jinja2 global autoescaping); "
  "Layer 6 — Server-side input validation (explicit NUMERIC_RANGES and VALID_* constant dicts "
  "for all 12 form fields). All 10 non-functional security requirements received PASS.")

h("5.3  System Features", 2)

for feat in [
    "Animated risk gauge (0–100 scale, colour-coded Green/Yellow/Orange/Red)",
    "Factor contributions breakdown — per-feature risk contribution (satisfies GDPR Article 22)",
    "Improvement roadmap — top-5 personalised recommendations with estimated score savings",
    "AI chatbot RiskBot: Claude Haiku + regex fallback for 20+ intent categories (EN and UZ)",
    "PDF report generation using ReportLab (average 1.1 seconds for a 3-page report)",
    "Bilingual interface: English / Uzbek toggle (200+ string pairs in lang.py)",
    "CSV export of full assessment history",
    "Insurance premium calculator (multipliers: 1.0× Low → 2.30× Very High)",
    "Driver comparison tool — side-by-side analysis of two profiles",
    "Progressive Web App (PWA) with service worker for offline static pages",
]:
    bullet(feat)

doc.add_paragraph()

# ═══════════════════════════════════════════════
# CHAPTER 6 — FINDINGS AND DISCUSSION
# ═══════════════════════════════════════════════
h("6.  Findings and Discussion")

p("All three research questions are answered with quantitative evidence. RQ1 (Can ML outperform "
  "traditional models?): Gradient Boosting achieves 91.3% accuracy versus 79.1% for Logistic "
  "Regression — a 12.2-point advantage. For a 10,000-policyholder portfolio, this equates to "
  "approximately 1,220 fewer misclassified drivers, or roughly £500,000 in correctly priced "
  "premiums annually at an average premium of £800.")

p("RQ2 (Which features are most important?): Previous accidents (26.5%) is by far the strongest "
  "predictor, consistent with Jiang et al. (2020) across 47 studies. Driving experience (17.8%), "
  "driver age (14.2%) and traffic violations (11.8%) complete the top four, together accounting "
  "for 69.6% of predictive power. Gender (0.6%) and marital status (1.1%) rank last — confirming "
  "these demographic features add almost no predictive value, supporting the case for removing "
  "them before regulatory submission.")

p("RQ3 (Can ML scores be made explainable for non-technical users?): The three-level design — "
  "animated gauge for quick impression, factor contributions for GDPR Article 22 compliance, and "
  "improvement roadmap for actionable dialogue — demonstrates that this is achievable. The AI "
  "chatbot, PDF export and model report page extend explainability to conversational, printed and "
  "regulatory channels.")

p("The primary limitation is the use of synthetic rather than real insurance claims data. While "
  "the synthetic dataset captures the key structural relationships from the literature, it cannot "
  "replicate seasonal variation, geographic clustering or vehicle model correlations present in a "
  "real portfolio. Validation on real claims data under a data sharing agreement with an insurance "
  "partner is the top priority for future work.")

# ═══════════════════════════════════════════════
# CHAPTER 7 — CONCLUSION AND RECOMMENDATIONS
# ═══════════════════════════════════════════════
h("7.  Conclusion and Recommendations")

p("RiskGuard AI met or exceeded all eight project objectives. Gradient Boosting achieved 91.3% "
  "accuracy, surpassing the 85% target by 6.3 percentage points. The deployed system provides "
  "three levels of explainability, a production-ready security architecture and a bilingual "
  "interface — demonstrating that ML-powered insurance underwriting is achievable using exclusively "
  "free, open-source tools within a ten-week development cycle.")

p("Five priority recommendations for future development: R1 — Integrate real telematics data (GPS "
  "speed profiles, braking events) to improve external validity. R2 — Implement TreeSHAP "
  "explainability to replace the current rule-based contribution layer, providing theoretically "
  "rigorous GDPR Article 22 compliance. R3 — Extend to a multi-tenant architecture (PostgreSQL, "
  "Docker, role-based access control) for enterprise deployment. R4 — Build an automated pytest "
  "suite with ≥80% branch coverage and GitHub Actions CI/CD. R5 — Evaluate LightGBM for "
  "production-scale portfolios (100,000+ policies) where training speed becomes significant.")

# ═══════════════════════════════════════════════
# CHAPTER 8 — PERSONAL DEVELOPMENT (GIBBS)
# ═══════════════════════════════════════════════
h("8.  Personal and Professional Development (Gibbs' Reflective Cycle)")

p("Applying Gibbs' six-stage cycle (1988) — Description, Feelings, Evaluation, Analysis, "
  "Conclusion, Action Plan — the most technically demanding moment was the Flask-Limiter / "
  "Flask-WTF conflict in Sprint 4. The default Redis storage backend required by Flask-Limiter "
  "was unavailable, causing an import error that blocked progress for two days. Switching to "
  "storage_uri='memory://' resolved the conflict — but only after reading both libraries' "
  "documentation in full. The lesson: always audit library dependencies before integration, "
  "not after an error appears.")

p("The most significant personal lesson was the gap between 'the code works' and 'I can explain "
  "every design decision clearly in academic prose'. Writing the report took substantially longer "
  "than estimated because design rationale had not been documented as decisions were made. In "
  "future projects I will maintain a decision log — one brief entry per significant architectural "
  "choice — which would both accelerate report writing and provide onboarding documentation for "
  "future team members. This lesson is captured in the action plan and the SWOT analysis of "
  "Chapter 7 of the full report.")

# ═══════════════════════════════════════════════
# CHAPTER 9 — PROJECT LINKS AND TECH STACK
# ═══════════════════════════════════════════════
h("9.  Project Resources and Technology Stack")

p("All project resources are open-source and publicly accessible. The links below are provided "
  "for demonstration during the viva examination:", size=12)

doc.add_paragraph()
link_line("Live Demonstration (Cloudflare Pages)", LIVE_URL)
link_line("Source Code (GitHub Repository)", GITHUB)
link_line("GitHub README (setup instructions)", GITHUB + "#readme",
          GITHUB + "#readme")
p("Local setup:  run.bat  (Windows) → http://127.0.0.1:5000    |    Login: admin / admin123",
  italic=True, size=11)

doc.add_paragraph()
tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = 'Table Grid'
tbl_hdr(tbl3, ["Technology", "Version", "Purpose"])
for row in [
    ("Python",           "3.11+",  "Primary backend language"),
    ("Flask",            "3.0",    "Web framework"),
    ("scikit-learn",     "1.4",    "ML pipeline and model training"),
    ("Gradient Boosting","sklearn","Production risk classifier"),
    ("SQLite",           "3.x",    "Assessment history database"),
    ("SQLAlchemy",       "2.0",    "ORM and parameterised queries"),
    ("Bootstrap",        "5.3",    "Responsive frontend CSS"),
    ("Chart.js",         "4.4",    "Dashboard trend charts"),
    ("ReportLab",        "4.x",    "PDF report generation"),
    ("Anthropic API",    "Haiku",  "AI chatbot backend"),
    ("Cloudflare Pages", "—",      "Free production hosting"),
    ("Flask-WTF",        "1.2",    "CSRF protection"),
    ("Flask-Limiter",    "3.x",    "Rate limiting (brute-force protection)"),
]:
    tbl_row(tbl3, row)
doc.add_paragraph()

# ═══════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════
h("References")

refs = [
    "ABI (2023) Motor Insurance Premium Tracker — Full Year 2023. London: Association of British Insurers.",
    "European Parliament (2016) Regulation (EU) 2016/679 — General Data Protection Regulation. Brussels: Official Journal of the EU.",
    "European Parliament (2024) Regulation (EU) 2024/1689 — Artificial Intelligence Act. Brussels: Official Journal of the EU.",
    "Gibbs, G. (1988) Learning by Doing: A Guide to Teaching and Learning Methods. Oxford: Oxford Polytechnic.",
    "Guelman, L. (2012) 'Gradient boosting trees for auto insurance loss cost modeling and prediction', Expert Systems with Applications, 39(3), pp.3659–3667.",
    "Hevner, A.R., March, S.T., Park, J. and Ram, S. (2004) 'Design science in information systems research', MIS Quarterly, 28(1), pp.75–105.",
    "Jiang, P., Fu, X. and Klemes, J.J. (2020) 'A comprehensive review of ML approaches for predicting road accident risks', Transportation Research Part C, 118, p.102760.",
    "Lundberg, S.M. and Lee, S.I. (2017) 'A unified approach to interpreting model predictions', Advances in Neural Information Processing Systems, 30, pp.4765–4774.",
    "McKinsey and Company (2023) Global Insurance Report 2023: Reimagining Insurance for the Digital Age. New York: McKinsey.",
    "Oates, B.J., Griffiths, M. and McLean, R. (2022) Researching Information Systems and Computing. 2nd edn. London: SAGE.",
    "Peffers, K., Tuunanen, T., Rothenberger, M.A. and Chatterjee, S. (2007) 'A design science research methodology for IS research', Journal of MIS, 24(3), pp.45–77.",
    "Pedregosa, F. et al. (2011) 'Scikit-learn: Machine learning in Python', Journal of Machine Learning Research, 12, pp.2825–2830.",
    "Saunders, M., Lewis, P. and Thornhill, A. (2023) Research Methods for Business Students. 9th edn. Harlow: Pearson.",
    "Uzbekistan Insurance Market Report (2024) Insurance Sector Development 2020–2024. Tashkent: Ministry of Finance.",
    "Verbelen, R., Antonio, K. and Claeskens, G. (2018) 'Unravelling the predictive power of telematics data in car insurance pricing', JRSS Series A, 181(4), pp.1055–1080.",
]
for ref in refs:
    pp = doc.add_paragraph(ref, style='List Bullet')
    if pp.runs:
        pp.runs[0].font.size = Pt(10)
        pp.runs[0].font.name = 'Times New Roman'
    pp.paragraph_format.space_after = Pt(4)

# ── Save ──────────────────────────────────────
doc.save(OUTPUT)
full_text = " ".join(pp.text for pp in doc.paragraphs)
wc = len(re.findall(r'\b\w+\b', full_text))
print(f"\n{'='*55}")
print(f"  Saved to:  {OUTPUT}")
print(f"  Live URL:  {LIVE_URL}")
print(f"  GitHub:    {GITHUB}")
print(f"  Word count (paragraphs): {wc:,}")
print(f"  Estimated total (incl. tables): ~{wc + 350:,}")
print(f"{'='*55}")
